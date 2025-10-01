"""
Document generation system with Word templates and PDF conversion
Supports template-based document creation with placeholders
"""

import os
import json
import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
try:
    from docx import Document
    from docxtpl import DocxTemplate
    DOCX_AVAILABLE = True
except ImportError as e:
    print(f"Warning: docx modules not available: {e}")
    print("Please install: pip install python-docx docxtpl")
    DOCX_AVAILABLE = False
    Document = None
    DocxTemplate = None
import subprocess
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from .templates_polifusion import polifusion_generator
from .templates_polifusion_enhanced import polifusion_enhanced_generator
from .templates_professional import ProfessionalTemplateGenerator
from .templates_ultra_professional import UltraProfessionalTemplateGenerator

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Document generation system with Word templates"""
    
    def __init__(self):
        # Use SHARED_DOCUMENTS_PATH if available, otherwise use MEDIA_ROOT
        base_path = getattr(settings, 'SHARED_DOCUMENTS_PATH', settings.MEDIA_ROOT)
        self.templates_path = os.path.join(base_path, 'templates')
        self.documents_path = os.path.join(base_path, 'documents')
        self.temp_path = os.path.join(base_path, 'temp')
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Ensure all required directories exist"""
        for path in [self.templates_path, self.documents_path, self.temp_path]:
            os.makedirs(path, exist_ok=True)
    
    def generate_document(self, template_name: str, incident_data: Dict, 
                         document_type: str = 'cliente_informe') -> Dict:
        """
        Generate a document from template with incident data
        
        Args:
            template_name: Name of the template to use
            incident_data: Dictionary with incident data
            document_type: Type of document (cliente_informe, proveedor_carta, lab_report)
            
        Returns:
            Dict with document paths and metadata
        """
        if not DOCX_AVAILABLE:
            raise ImportError("Document generation requires python-docx and docxtpl packages")
        try:
            # Get template path
            template_path = os.path.join(self.templates_path, f"{template_name}.docx")
            
            if not os.path.exists(template_path):
                # Create default template if it doesn't exist
                template_path = self._create_default_template(template_name, document_type)
            
            # Load template
            template = DocxTemplate(template_path)
            
            # Prepare context data
            context = self._prepare_context(incident_data, document_type)
            
            # Render template
            template.render(context)
            
            # Generate document filename
            doc_filename = self._generate_filename(incident_data, document_type)
            doc_path = os.path.join(self.documents_path, doc_filename)
            
            # Save document
            template.save(doc_path)
            
            # Convert to PDF
            pdf_path = self._convert_to_pdf(doc_path)
            
            # Generate metadata
            metadata = self._generate_metadata(incident_data, document_type, doc_path, pdf_path)
            
            return {
                'success': True,
                'docx_path': doc_path,
                'pdf_path': pdf_path,
                'filename': doc_filename,
                'metadata': metadata,
                'template_used': template_name,
                'context_used': context
            }
            
        except Exception as e:
            logger.error(f"Error generating document: {e}")
            return {
                'success': False,
                'error': str(e),
                'template_name': template_name,
                'document_type': document_type
            }
    
    def _create_default_template(self, template_name: str, document_type: str) -> str:
        """Create a default template if it doesn't exist"""
        try:
            template_path = os.path.join(self.templates_path, f"{template_name}.docx")
            
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'Plantilla {template_name}', 0)
            
            # Add content based on document type
            if document_type == 'cliente_informe':
                self._add_client_report_content(doc)
            elif document_type == 'proveedor_carta':
                self._add_supplier_letter_content(doc)
            elif document_type == 'lab_report':
                self._add_lab_report_content(doc)
            else:
                self._add_generic_content(doc)
            
            # Save template
            doc.save(template_path)
            logger.info(f"Created default template: {template_path}")
            
            return template_path
            
        except Exception as e:
            logger.error(f"Error creating default template: {e}")
            raise
    
    def _add_client_report_content(self, doc: Document):
        """Add content for client report template"""
        doc.add_heading('Informe de Incidencia de Postventa', level=1)
        
        # Add placeholders
        doc.add_paragraph('Cliente: {{CLIENTE}}')
        doc.add_paragraph('Fecha de Detección: {{FECHA_DETECCION}}')
        doc.add_paragraph('SKU: {{SKU}}')
        doc.add_paragraph('Lote: {{LOTE}}')
        doc.add_paragraph('')
        
        doc.add_heading('Descripción del Problema', level=2)
        doc.add_paragraph('{{DESCRIPCION}}')
        doc.add_paragraph('')
        
        doc.add_heading('Conclusiones Técnicas', level=2)
        doc.add_paragraph('{{CONCLUSIONES_TECNICAS}}')
        doc.add_paragraph('')
        
        doc.add_heading('Recomendaciones', level=2)
        doc.add_paragraph('{{RECOMENDACIONES}}')
        doc.add_paragraph('')
        
        doc.add_paragraph('Firmante: {{FIRMANTE}}')
        doc.add_paragraph('Fecha: {{FECHA_FIRMA}}')
    
    def _add_supplier_letter_content(self, doc: Document):
        """Add content for supplier letter template"""
        doc.add_heading('Carta Técnica a Proveedor', level=1)
        
        # Add placeholders
        doc.add_paragraph('Proveedor: {{PROVEEDOR}}')
        doc.add_paragraph('Lote: {{LOTE}}')
        doc.add_paragraph('Número de Pedido: {{NUM_PEDIDO}}')
        doc.add_paragraph('')
        
        doc.add_heading('Descripción del Problema', level=2)
        doc.add_paragraph('{{DESCRIPCION}}')
        doc.add_paragraph('')
        
        doc.add_heading('Análisis Técnico', level=2)
        doc.add_paragraph('{{ANALISIS_TECNICO}}')
        doc.add_paragraph('')
        
        doc.add_heading('Conclusiones Técnicas', level=2)
        doc.add_paragraph('{{CONCLUSIONES_TECNICAS}}')
        doc.add_paragraph('')
        
        doc.add_heading('Acciones Requeridas', level=2)
        doc.add_paragraph('{{ACCIONES_REQUERIDAS}}')
        doc.add_paragraph('')
        
        doc.add_paragraph('Firmante: {{FIRMANTE}}')
        doc.add_paragraph('Fecha: {{FECHA_FIRMA}}')
    
    def _add_lab_report_content(self, doc: Document):
        """Add content for lab report template"""
        doc.add_heading('Reporte de Laboratorio', level=1)
        
        # Add placeholders
        doc.add_paragraph('Incidente: {{INCIDENTE}}')
        doc.add_paragraph('Muestra: {{MUESTRA}}')
        doc.add_paragraph('Fecha de Análisis: {{FECHA_ANALISIS}}')
        doc.add_paragraph('')
        
        doc.add_heading('Ensayos Realizados', level=2)
        doc.add_paragraph('{{ENSAYOS}}')
        doc.add_paragraph('')
        
        doc.add_heading('Observaciones', level=2)
        doc.add_paragraph('{{OBSERVACIONES}}')
        doc.add_paragraph('')
        
        doc.add_heading('Resultados', level=2)
        doc.add_paragraph('{{RESULTADOS}}')
        doc.add_paragraph('')
        
        doc.add_heading('Conclusiones', level=2)
        doc.add_paragraph('{{CONCLUSIONES}}')
        doc.add_paragraph('')
        
        doc.add_paragraph('Experto: {{EXPERTO}}')
        doc.add_paragraph('Fecha: {{FECHA_FIRMA}}')
    
    def _add_generic_content(self, doc: Document):
        """Add generic content for unknown document types"""
        doc.add_heading('Documento Generado', level=1)
        doc.add_paragraph('{{CONTENIDO}}')
        doc.add_paragraph('')
        doc.add_paragraph('Fecha: {{FECHA}}')
        doc.add_paragraph('Autor: {{AUTOR}}')
    
    def _prepare_context(self, incident_data: Dict, document_type: str) -> Dict:
        """Prepare context data for template rendering"""
        context = {
            # Basic incident data
            'CLIENTE': incident_data.get('cliente', 'No especificado'),
            'PROVEEDOR': incident_data.get('proveedor', 'No especificado'),
            'SKU': incident_data.get('sku', 'No especificado'),
            'LOTE': incident_data.get('lote', 'No especificado'),
            'NUM_PEDIDO': incident_data.get('num_pedido', 'No especificado'),
            'FECHA_DETECCION': incident_data.get('fecha_deteccion', 'No especificada'),
            'DESCRIPCION': incident_data.get('descripcion', 'No especificada'),
            'FECHA_FIRMA': incident_data.get('fecha_firma', 'No especificada'),
            'FIRMANTE': incident_data.get('firmante', 'No especificado'),
            
            # Technical data
            'CONCLUSIONES_TECNICAS': incident_data.get('conclusiones_tecnicas', 'No especificadas'),
            'RECOMENDACIONES': incident_data.get('recomendaciones', 'No especificadas'),
            'ANALISIS_TECNICO': incident_data.get('analisis_tecnico', 'No especificado'),
            'ACCIONES_REQUERIDAS': incident_data.get('acciones_requeridas', 'No especificadas'),
            
            # Lab report data
            'INCIDENTE': incident_data.get('incidente', 'No especificado'),
            'MUESTRA': incident_data.get('muestra', 'No especificada'),
            'ENSAYOS': incident_data.get('ensayos', 'No especificados'),
            'OBSERVACIONES': incident_data.get('observaciones', 'No especificadas'),
            'RESULTADOS': incident_data.get('resultados', 'No especificados'),
            'CONCLUSIONES': incident_data.get('conclusiones', 'No especificadas'),
            'EXPERTO': incident_data.get('experto', 'No especificado'),
            
            # Generic data
            'CONTENIDO': incident_data.get('contenido', 'No especificado'),
            'FECHA': incident_data.get('fecha', 'No especificada'),
            'AUTOR': incident_data.get('autor', 'No especificado'),
        }
        
        # Add AI-generated content if available
        if 'ai_analysis' in incident_data:
            ai_data = incident_data['ai_analysis']
            context.update({
                'CONCLUSIONES_TECNICAS': ai_data.get('conclusiones', context['CONCLUSIONES_TECNICAS']),
                'RECOMENDACIONES': ai_data.get('recomendaciones', context['RECOMENDACIONES']),
                'ANALISIS_TECNICO': ai_data.get('analisis', context['ANALISIS_TECNICO']),
            })
        
        return context
    
    def _generate_filename(self, incident_data: Dict, document_type: str) -> str:
        """Generate filename for the document"""
        import datetime
        
        # Get current date
        now = datetime.datetime.now()
        date_str = now.strftime("%Y%m%d")
        
        # Get incident info
        cliente = incident_data.get('cliente', 'cliente').replace(' ', '_')
        sku = incident_data.get('sku', 'sku')
        lote = incident_data.get('lote', 'lote')
        
        # Generate filename
        filename = f"{document_type}_{cliente}_{sku}_{lote}_{date_str}.docx"
        
        return filename
    
    def _convert_to_pdf(self, docx_path: str) -> str:
        """Convert DOCX to PDF using LibreOffice"""
        try:
            # Generate PDF path
            pdf_path = docx_path.replace('.docx', '.pdf')
            
            # Use LibreOffice to convert
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path),
                docx_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and os.path.exists(pdf_path):
                logger.info(f"Successfully converted {docx_path} to {pdf_path}")
                return pdf_path
            else:
                logger.error(f"Error converting to PDF: {result.stderr}")
                return None
                
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            return None
    
    def _generate_metadata(self, incident_data: Dict, document_type: str, 
                          docx_path: str, pdf_path: str) -> Dict:
        """Generate metadata for the document"""
        import datetime
        
        metadata = {
            'document_type': document_type,
            'incident_id': incident_data.get('id'),
            'cliente': incident_data.get('cliente'),
            'sku': incident_data.get('sku'),
            'lote': incident_data.get('lote'),
            'created_at': datetime.datetime.now().isoformat(),
            'docx_path': docx_path,
            'pdf_path': pdf_path,
            'docx_size': os.path.getsize(docx_path) if os.path.exists(docx_path) else 0,
            'pdf_size': os.path.getsize(pdf_path) if pdf_path and os.path.exists(pdf_path) else 0,
            'template_used': f"{document_type}.docx",
            'ai_generated': 'ai_analysis' in incident_data
        }
        
        return metadata
    
    def batch_generate_documents(self, incidents: List[Dict], document_type: str) -> List[Dict]:
        """Generate documents for multiple incidents"""
        results = []
        for incident in incidents:
            result = self.generate_document(
                template_name=document_type,
                incident_data=incident,
                document_type=document_type
            )
            results.append(result)
        return results
    
    def get_available_templates(self) -> List[str]:
        """Get list of available templates"""
        try:
            templates = []
            for file in os.listdir(self.templates_path):
                if file.endswith('.docx'):
                    templates.append(file.replace('.docx', ''))
            return templates
        except Exception as e:
            logger.error(f"Error getting available templates: {e}")
            return []
    
    def create_custom_template(self, template_name: str, content: str) -> bool:
        """Create a custom template from content"""
        try:
            template_path = os.path.join(self.templates_path, f"{template_name}.docx")
            
            # Create document
            doc = Document()
            
            # Add content
            doc.add_paragraph(content)
            
            # Save template
            doc.save(template_path)
            
            logger.info(f"Created custom template: {template_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating custom template: {e}")
            return False
    
    def get_document_statistics(self) -> Dict:
        """Get statistics about generated documents"""
        try:
            stats = {
                'total_documents': 0,
                'total_size': 0,
                'documents_by_type': {},
                'recent_documents': []
            }
            
            # Count documents
            for file in os.listdir(self.documents_path):
                if file.endswith('.docx'):
                    stats['total_documents'] += 1
                    file_path = os.path.join(self.documents_path, file)
                    stats['total_size'] += os.path.getsize(file_path)
                    
                    # Categorize by type
                    doc_type = file.split('_')[0]
                    stats['documents_by_type'][doc_type] = stats['documents_by_type'].get(doc_type, 0) + 1
            
            return stats
            
        except Exception as e:
            logger.error(f"Error getting document statistics: {e}")
            return {}
    
    def generate_polifusion_lab_report(self, incident_data: Dict) -> Dict:
        """Generate Polifusión laboratory report with ultra professional template"""
        try:
            # Create ultra professional template if it doesn't exist
            template_path = os.path.join(self.templates_path, "polifusion_lab_report_ultra_professional.docx")
            if not os.path.exists(template_path):
                ultra_generator = UltraProfessionalTemplateGenerator()
                ultra_generator.create_ultra_professional_lab_report(template_path)
                logger.info(f"Created ultra professional Polifusión lab report template: {template_path}")
            
            # Load template
            template = DocxTemplate(template_path)
            
            # Prepare ultra professional Polifusión-specific context
            context = self._prepare_polifusion_ultra_professional_context(incident_data)
            
            # Render template
            template.render(context)
            
            # Generate document filename
            doc_filename = self._generate_polifusion_filename(incident_data)
            doc_path = os.path.join(self.documents_path, doc_filename)
            
            # Save document
            template.save(doc_path)
            
            # Convert to PDF
            pdf_path = self._convert_to_pdf(doc_path)
            
            return {
                'success': True,
                'docx_path': doc_path,
                'pdf_path': pdf_path,
                'filename': doc_filename,
                'template_used': 'polifusion_lab_report_ultra_professional',
                'context_used': context
            }
            
        except Exception as e:
            logger.error(f"Error generating ultra professional Polifusión lab report: {e}")
            return {
                'success': False,
                'error': str(e),
                'template_name': 'polifusion_lab_report_ultra_professional'
            }
    
    def _prepare_polifusion_context(self, incident_data: Dict) -> Dict:
        """Prepare context data for Polifusión templates"""
        context = {
            # Información del solicitante
            'SOLICITANTE': incident_data.get('solicitante', 'POLIFUSION'),
            'FECHA_SOLICITUD': incident_data.get('fecha_solicitud', datetime.now().strftime('%d/%m/%Y')),
            'CLIENTE': incident_data.get('cliente', 'POLIFUSION'),
            
            # Información técnica
            'DIAMETRO': incident_data.get('diametro', '160'),
            'PROYECTO': incident_data.get('proyecto', 'Proyecto Alameda Park'),
            'UBICACION': incident_data.get('ubicacion', 'Av. Libertador Bernardo O\'Higgins 4687'),
            'PRESION': incident_data.get('presion', '11.8-12'),
            'TEMPERATURA': incident_data.get('temperatura', 'No registrada'),
            'INFORMANTE': incident_data.get('informante', 'Yenny Valdivia Sazo'),
            
            # Ensayos
            'ENSAYOS_ADICIONALES': incident_data.get('ensayos_adicionales', 'Análisis de fractura y cristalización'),
            
            # Comentarios detallados
            'COMENTARIOS_DETALLADOS': incident_data.get('comentarios_detallados', 
                'Se recibió una muestra compuesta por un codo de 90° PP-RCT de 160mm, fusionado en ambos extremos a una tubería PP-RCT/FIBERGLASS S-3.2 de 160mm. La muestra presenta una grieta longitudinal tipo corte en la sección transversal de la tubería en uno de los extremos del codo, con un inserto de tapón, posiblemente para detener una fuga.'),
            
            # Conclusiones
            'CONCLUSIONES_DETALLADAS': incident_data.get('conclusiones_detalladas',
                'El análisis de la muestra reveló que los extremos de la tubería no fueron reducidos, y el cordón de fusión externo era irregular, mostrando fusión excesiva y evidencia clara de sobre-fusión de la pared del fitting. También se observó una fisura tipo corte o grieta longitudinal, en la cual se insertó un tapón de reparación, presumiblemente como intento de contener la fuga.'),
            
            # Experto
            'EXPERTO_NOMBRE': incident_data.get('experto_nombre', 'CÉSAR MUNIZAGA GARRIDO'),
            
            # Análisis detallado
            'ANALISIS_DETALLADO': incident_data.get('analisis_detallado',
                'Se realizó un corte longitudinal del fitting, revelando cordones de fusión internos corrugados, lo cual es una indicación clara de presencia de agua durante el proceso de fusión. La inspección de la grieta, encontrada en el punto donde se inserta una sección de tubería, justo detrás del cordón de fusión, reveló una separación de 1-2 mm de la pared del fitting. Al aplicar una fuerza de flexión se reveló una fractura frágil característica, típica de material cristalizado debido al calentamiento excesivo.')
        }
        
        return context
    
    def _prepare_polifusion_ultra_professional_context(self, incident_data: Dict) -> Dict:
        """Prepare ultra professional context data for Polifusión lab templates"""
        now = datetime.now()
        
        context = {
            # Información del solicitante
            'SOLICITANTE': incident_data.get('solicitante', 'POLIFUSIÓN'),
            'FECHA_SOLICITUD': incident_data.get('fecha_solicitud', now.strftime('%d/%m/%Y')),
            'CLIENTE': incident_data.get('cliente', 'POLIFUSIÓN'),
            'INFORMANTE': incident_data.get('informante', ''),
            
            # Información técnica
            'DIAMETRO': incident_data.get('diametro', '160'),
            'PROYECTO': incident_data.get('proyecto', ''),
            'UBICACION': incident_data.get('ubicacion', ''),
            'PRESION': incident_data.get('presion', ''),
            'TEMPERATURA': incident_data.get('temperatura', 'No registrada'),
            
            # Ensayos
            'ENSAYOS_ADICIONALES': incident_data.get('ensayos_adicionales', ''),
            
            # Comentarios detallados
            'COMENTARIOS_DETALLADOS': incident_data.get('comentarios_detallados', ''),
            
            # Conclusiones
            'CONCLUSIONES_DETALLADAS': incident_data.get('conclusiones_detalladas', ''),
            
            # Experto
            'EXPERTO_NOMBRE': incident_data.get('experto_nombre', ''),
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context

    def _prepare_polifusion_incident_ultra_professional_context(self, incident_data: Dict) -> Dict:
        """Prepare ultra professional context data for Polifusión incident templates"""
        now = datetime.now()
        
        context = {
            # Información de la incidencia
            'PROVEEDOR': incident_data.get('proveedor', ''),
            'OBRA': incident_data.get('obra', ''),
            'CLIENTE': incident_data.get('cliente', ''),
            'FECHA_INCIDENCIA': incident_data.get('fecha_deteccion', now.strftime('%d/%m/%Y')),
            
            # Descripción del problema
            'DESCRIPCION_PROBLEMA': incident_data.get('descripcion_problema', ''),
            
            # Acciones inmediatas
            'ACCIONES_INMEDIATAS': incident_data.get('acciones_inmediatas', ''),
            
            # Evolución y acciones posteriores
            'EVOLUCION_ACCIONES': incident_data.get('evolucion_acciones', ''),
            
            # Observaciones y cierre
            'OBSERVACIONES': incident_data.get('observaciones', ''),
            'OBSERVACIONES_CIERRE': incident_data.get('observaciones', ''),
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context

    def _prepare_polifusion_visit_ultra_professional_context(self, visit_data: Dict) -> Dict:
        """Prepare ultra professional context data for Polifusión visit templates"""
        now = datetime.now()
        
        context = {
            # Información de la visita
            'OBRA': visit_data.get('obra', ''),
            'CLIENTE': visit_data.get('cliente', ''),
            'VENDEDOR': visit_data.get('vendedor', ''),
            'TECNICO': visit_data.get('tecnico', ''),
            'FECHA_VISITA': visit_data.get('fecha_visita', now.strftime('%d/%m/%Y')),
            
            # Personal presente
            'PERSONAL_INFO': visit_data.get('personal_info', ''),
            'ROLES_CONTACTOS': visit_data.get('roles_contactos', ''),
            
            # Evaluación técnica
            'MAQUINARIA_USO': visit_data.get('maquinaria_uso', ''),
            
            # Observaciones
            'OBSERVACIONES_INSTALACION': visit_data.get('observaciones_instalacion', ''),
            'OBSERVACIONES_MATERIAL': visit_data.get('observaciones_material', ''),
            'OBSERVACIONES_TECNICO': visit_data.get('observaciones_tecnico', ''),
            'OBSERVACIONES_GENERAL': visit_data.get('observaciones_general', ''),
            
            # Recomendaciones
            'RECOMENDACIONES': visit_data.get('recomendaciones', ''),
            
            # Firmas
            'FIRMA_VENDEDOR': visit_data.get('firma_vendedor', ''),
            'FIRMA_TECNICO': visit_data.get('firma_tecnico', ''),
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context

    def _prepare_polifusion_professional_context(self, incident_data: Dict) -> Dict:
        """Prepare professional context data for Polifusión lab templates"""
        now = datetime.now()
        
        context = {
            # Información del solicitante
            'SOLICITANTE': incident_data.get('solicitante', 'POLIFUSIÓN'),
            'FECHA_SOLICITUD': incident_data.get('fecha_solicitud', now.strftime('%d/%m/%Y')),
            'CLIENTE': incident_data.get('cliente', 'POLIFUSIÓN'),
            'INFORMANTE': incident_data.get('informante', ''),
            
            # Información técnica
            'DIAMETRO': incident_data.get('diametro', '160'),
            'PROYECTO': incident_data.get('proyecto', ''),
            'UBICACION': incident_data.get('ubicacion', ''),
            'PRESION': incident_data.get('presion', ''),
            'TEMPERATURA': incident_data.get('temperatura', 'No registrada'),
            
            # Ensayos
            'ENSAYOS_ADICIONALES': incident_data.get('ensayos_adicionales', ''),
            
            # Comentarios detallados
            'COMENTARIOS_DETALLADOS': incident_data.get('comentarios_detallados', ''),
            
            # Conclusiones
            'CONCLUSIONES_DETALLADAS': incident_data.get('conclusiones_detalladas', ''),
            
            # Experto
            'EXPERTO_NOMBRE': incident_data.get('experto_nombre', ''),
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context

    def _prepare_polifusion_enhanced_context(self, incident_data: Dict) -> Dict:
        """Prepare enhanced context data for Polifusión templates"""
        now = datetime.now()
        
        context = {
            # Información del solicitante
            'SOLICITANTE': incident_data.get('solicitante', 'POLIFUSIÓN'),
            'FECHA_SOLICITUD': incident_data.get('fecha_solicitud', now.strftime('%d/%m/%Y')),
            'CLIENTE': incident_data.get('cliente', 'POLIFUSIÓN'),
            'INFORMANTE': incident_data.get('informante', 'Yenny Valdivia Sazo'),
            
            # Información técnica
            'DIAMETRO': incident_data.get('diametro', '160'),
            'PROYECTO': incident_data.get('proyecto', 'Proyecto Alameda Park'),
            'UBICACION': incident_data.get('ubicacion', 'Av. Libertador Bernardo O\'Higgins 4687'),
            'PRESION': incident_data.get('presion', '11.8-12'),
            'TEMPERATURA': incident_data.get('temperatura', 'No registrada'),
            
            # Ensayos
            'ENSAYOS_ADICIONALES': incident_data.get('ensayos_adicionales', 'Análisis de fractura y cristalización'),
            
            # Comentarios detallados
            'COMENTARIOS_DETALLADOS': incident_data.get('comentarios_detallados', 'Se recibió una muestra compuesta por un codo de 90° PP-RCT de 160mm, fusionado en ambos extremos a una tubería PP-RCT/FIBERGLASS S-3.2 de 160mm. La muestra presenta una grieta longitudinal tipo corte en la sección transversal de la tubería en uno de los extremos del codo, con un inserto de tapón, posiblemente para detener una fuga.'),
            
            # Conclusiones
            'CONCLUSIONES_DETALLADAS': incident_data.get('conclusiones_detalladas', 'El análisis de la muestra reveló que los extremos de la tubería no fueron reducidos, y el cordón de fusión externo era irregular, mostrando fusión excesiva y evidencia clara de sobre-fusión de la pared del fitting. También se observó una fisura tipo corte o grieta longitudinal, en la cual se insertó un tapón de reparación, presumiblemente como intento de contener la fuga.'),
            
            # Experto
            'EXPERTO_NOMBRE': incident_data.get('experto_nombre', 'CÉSAR MUNIZAGA GARRIDO'),
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context
    
    def _prepare_polifusion_incident_professional_context(self, incident_data: Dict) -> Dict:
        """Prepare professional context data for Polifusión incident templates"""
        now = datetime.now()
        
        context = {
            # Información de la incidencia
            'PROVEEDOR': incident_data.get('proveedor', ''),
            'OBRA': incident_data.get('obra', ''),
            'CLIENTE': incident_data.get('cliente', ''),
            'FECHA_INCIDENCIA': incident_data.get('fecha_deteccion', now.strftime('%d/%m/%Y')),
            
            # Descripción del problema
            'DESCRIPCION_PROBLEMA': incident_data.get('descripcion_problema', ''),
            
            # Acciones inmediatas
            'ACCIONES_INMEDIATAS': incident_data.get('acciones_inmediatas', ''),
            
            # Evolución y acciones posteriores
            'EVOLUCION_ACCIONES': incident_data.get('evolucion_acciones', ''),
            
            # Observaciones y cierre
            'OBSERVACIONES': incident_data.get('observaciones', ''),
            'OBSERVACIONES_CIERRE': incident_data.get('observaciones', ''),
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context

    def _prepare_polifusion_incident_enhanced_context(self, incident_data: Dict) -> Dict:
        """Prepare enhanced context data for Polifusión incident templates"""
        now = datetime.now()
        
        context = {
            # Información de la incidencia
            'PROVEEDOR': incident_data.get('proveedor', ''),
            'OBRA': incident_data.get('obra', ''),
            'CLIENTE': incident_data.get('cliente', ''),
            'FECHA_INCIDENCIA': incident_data.get('fecha_deteccion', now.strftime('%d/%m/%Y')),
            
            # Descripción del problema
            'DESCRIPCION_PROBLEMA': incident_data.get('descripcion_problema', ''),
            
            # Acciones inmediatas
            'ACCIONES_INMEDIATAS': incident_data.get('acciones_inmediatas', ''),
            
            # Evolución y acciones posteriores
            'EVOLUCION_ACCIONES': incident_data.get('evolucion_acciones', ''),
            
            # Observaciones y cierre
            'OBSERVACIONES_CIERRE': incident_data.get('observaciones', ''),
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context
    
    def _generate_polifusion_filename(self, incident_data: Dict) -> str:
        """Generate filename for Polifusión documents"""
        now = datetime.now()
        date_str = now.strftime("%Y%m%d")
        
        cliente = incident_data.get('cliente', 'polifusion').replace(' ', '_')
        proyecto = incident_data.get('proyecto', 'proyecto').replace(' ', '_')
        
        filename = f"INFORME_LAB_POLIFUSION_{cliente}_{proyecto}_{date_str}.docx"
        
        return filename
    
    def generate_polifusion_incident_report(self, incident_data: Dict) -> Dict:
        """Generate Polifusión incident report with ultra professional template"""
        try:
            # Create ultra professional template if it doesn't exist
            template_path = os.path.join(self.templates_path, "polifusion_incident_report_ultra_professional.docx")
            if not os.path.exists(template_path):
                ultra_generator = UltraProfessionalTemplateGenerator()
                ultra_generator.create_ultra_professional_incident_report(template_path)
                logger.info(f"Created ultra professional Polifusión incident report template: {template_path}")
            
            # Load template
            template = DocxTemplate(template_path)
            
            # Prepare ultra professional Polifusión-specific context
            context = self._prepare_polifusion_incident_ultra_professional_context(incident_data)
            
            # Render template
            template.render(context)
            
            # Generate document filename
            doc_filename = self._generate_polifusion_incident_filename(incident_data)
            doc_path = os.path.join(self.documents_path, doc_filename)
            
            # Save document
            template.save(doc_path)
            
            # Convert to PDF
            pdf_path = self._convert_to_pdf(doc_path)
            
            return {
                'success': True,
                'docx_path': doc_path,
                'pdf_path': pdf_path,
                'filename': doc_filename,
                'template_used': 'polifusion_incident_report_ultra_professional',
                'context_used': context
            }
            
        except Exception as e:
            logger.error(f"Error generating ultra professional Polifusión incident report: {e}")
            return {
                'success': False,
                'error': str(e),
                'template_name': 'polifusion_incident_report_ultra_professional'
            }
    
    def generate_polifusion_visit_report(self, visit_data: Dict) -> Dict:
        """Generate Polifusión visit report with ultra professional template"""
        try:
            # Create ultra professional template if it doesn't exist
            template_path = os.path.join(self.templates_path, "polifusion_visit_report_ultra_professional.docx")
            if not os.path.exists(template_path):
                ultra_generator = UltraProfessionalTemplateGenerator()
                ultra_generator.create_ultra_professional_visit_report(template_path)
                logger.info(f"Created ultra professional Polifusión visit report template: {template_path}")
            
            # Load template
            template = DocxTemplate(template_path)
            
            # Prepare ultra professional Polifusión-specific context
            context = self._prepare_polifusion_visit_ultra_professional_context(visit_data)
            
            # Render template
            template.render(context)
            
            # Generate document filename
            doc_filename = self._generate_polifusion_visit_filename(visit_data)
            doc_path = os.path.join(self.documents_path, doc_filename)
            
            # Save document
            template.save(doc_path)
            
            # Convert to PDF
            pdf_path = self._convert_to_pdf(doc_path)
            
            return {
                'success': True,
                'docx_path': doc_path,
                'pdf_path': pdf_path,
                'filename': doc_filename,
                'template_used': 'polifusion_visit_report_ultra_professional',
                'context_used': context
            }
            
        except Exception as e:
            logger.error(f"Error generating ultra professional Polifusión visit report: {e}")
            return {
                'success': False,
                'error': str(e),
                'template_name': 'polifusion_visit_report_ultra_professional'
            }
    
    def _prepare_polifusion_visit_professional_context(self, visit_data: Dict) -> Dict:
        """Prepare professional context data for Polifusión visit templates"""
        now = datetime.now()
        
        context = {
            # Información de la visita
            'OBRA': visit_data.get('obra', ''),
            'CLIENTE': visit_data.get('cliente', ''),
            'VENDEDOR': visit_data.get('vendedor', ''),
            'TECNICO': visit_data.get('tecnico', ''),
            'FECHA_VISITA': visit_data.get('fecha_visita', now.strftime('%d/%m/%Y')),
            
            # Personal presente
            'PERSONAL_INFO': visit_data.get('personal_info', ''),
            'ROLES_CONTACTOS': visit_data.get('roles_contactos', ''),
            
            # Evaluación técnica
            'MAQUINARIA_USO': visit_data.get('maquinaria_uso', ''),
            
            # Observaciones
            'OBSERVACIONES_INSTALACION': visit_data.get('observaciones_instalacion', ''),
            'OBSERVACIONES_MATERIAL': visit_data.get('observaciones_material', ''),
            'OBSERVACIONES_TECNICO': visit_data.get('observaciones_tecnico', ''),
            'OBSERVACIONES_GENERAL': visit_data.get('observaciones_general', ''),
            
            # Recomendaciones
            'RECOMENDACIONES': visit_data.get('recomendaciones', ''),
            
            # Firmas
            'FIRMA_VENDEDOR': visit_data.get('firma_vendedor', ''),
            'FIRMA_TECNICO': visit_data.get('firma_tecnico', ''),
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context

    def _prepare_polifusion_visit_enhanced_context(self, visit_data: Dict) -> Dict:
        """Prepare enhanced context data for Polifusión visit templates"""
        now = datetime.now()
        
        context = {
            # Información de la obra
            'OBRA': visit_data.get('obra', ''),
            'CLIENTE': visit_data.get('cliente', ''),
            'UBICACION': visit_data.get('direccion', ''),
            
            # Información del personal
            'VENDEDOR': visit_data.get('vendedor', ''),
            'TECNICO': visit_data.get('tecnico', ''),
            'FECHA_VISITA': visit_data.get('fecha_visita', now.strftime('%d/%m/%Y')),
            
            # Roles y contactos
            'ROLES_CONTACTOS': visit_data.get('encargado_calidad', '') + ' ' + visit_data.get('profesional_obra', ''),
            
            # Uso de maquinaria
            'USO_MAQUINARIA': visit_data.get('maquinaria', ''),
            
            # Observaciones por categoría
            'OBSERVACIONES_CATEGORIA': f"Muro/Tabique: {visit_data.get('obs_muro_tabique', '')}\nMatriz: {visit_data.get('obs_matriz', '')}\nLoza: {visit_data.get('obs_loza', '')}\nAlmacenaje: {visit_data.get('obs_almacenaje', '')}\nPre-armados: {visit_data.get('obs_pre_armados', '')}\nExteriores: {visit_data.get('obs_exteriores', '')}\nGenerales: {visit_data.get('obs_generales', '')}",
            
            # Fecha de generación
            'FECHA_GENERACION': now.strftime('%d/%m/%Y %H:%M')
        }
        
        return context
    
    def _prepare_polifusion_incident_context(self, incident_data: Dict) -> Dict:
        """Prepare context data for Polifusión incident templates"""
        context = {
            # Información del informe
            'REPORT_NUMBER': incident_data.get('report_number', ''),
            'REPORT_DATE': incident_data.get('report_date', datetime.now().strftime('%d/%m/%Y')),
            
            # Registro de información
            'PROVEEDOR': incident_data.get('proveedor', ''),
            'OBRA': incident_data.get('obra', ''),
            'PRODUCCION': incident_data.get('produccion', ''),
            'CLIENTE': incident_data.get('cliente', ''),
            'SERVICIO': incident_data.get('servicio', ''),
            'RUT': incident_data.get('rut', ''),
            'DIRECCION': incident_data.get('direccion', ''),
            'OTROS': incident_data.get('otros', ''),
            'CONTACTOS': incident_data.get('contactos', ''),
            'FECHA_DETECCION': incident_data.get('fecha_deteccion', ''),
            'HORA': incident_data.get('hora', ''),
            
            # Descripción del problema
            'DESCRIPCION_PROBLEMA': incident_data.get('descripcion_problema', ''),
            
            # Acciones inmediatas
            'ACCIONES_INMEDIATAS': incident_data.get('acciones_inmediatas', ''),
            
            # Evolución/Acciones posteriores
            'FECHA_1': incident_data.get('evolucion_acciones', [{}])[0].get('fecha', '') if incident_data.get('evolucion_acciones') else '',
            'DESCRIPCION_1': incident_data.get('evolucion_acciones', [{}])[0].get('descripcion', '') if incident_data.get('evolucion_acciones') else '',
            'FECHA_2': incident_data.get('evolucion_acciones', [{}])[1].get('fecha', '') if len(incident_data.get('evolucion_acciones', [])) > 1 else '',
            'DESCRIPCION_2': incident_data.get('evolucion_acciones', [{}])[1].get('descripcion', '') if len(incident_data.get('evolucion_acciones', [])) > 1 else '',
            'FECHA_3': incident_data.get('evolucion_acciones', [{}])[2].get('fecha', '') if len(incident_data.get('evolucion_acciones', [])) > 2 else '',
            'DESCRIPCION_3': incident_data.get('evolucion_acciones', [{}])[2].get('descripcion', '') if len(incident_data.get('evolucion_acciones', [])) > 2 else '',
            
            # Observaciones y cierre
            'OBSERVACIONES': incident_data.get('observaciones', ''),
            'FECHA_CIERRE': incident_data.get('fecha_cierre', ''),
        }
        
        return context
    
    def _prepare_polifusion_visit_context(self, visit_data: Dict) -> Dict:
        """Prepare context data for Polifusión visit templates"""
        context = {
            # Información del reporte
            'ORDEN_NUMBER': visit_data.get('orden_number', ''),
            'FECHA_VISITA': visit_data.get('fecha_visita', datetime.now().strftime('%d/%m/%Y')),
            
            # Información general de la obra
            'OBRA': visit_data.get('obra', ''),
            'CLIENTE': visit_data.get('cliente', ''),
            'DIRECCION': visit_data.get('direccion', ''),
            'ADMINISTRADOR': visit_data.get('administrador', ''),
            'CONSTRUCTOR': visit_data.get('constructor', ''),
            'MOTIVO_VISITA': visit_data.get('motivo_visita', '01-Visita Técnica'),
            
            # Información del personal
            'VENDEDOR': visit_data.get('vendedor', ''),
            'COMUNA': visit_data.get('comuna', ''),
            'CIUDAD': visit_data.get('ciudad', ''),
            'INSTALADOR': visit_data.get('instalador', ''),
            'FONO_INSTALADOR': visit_data.get('fono_instalador', ''),
            'TECNICO': visit_data.get('tecnico', ''),
            
            # Roles y contactos
            'ENCARGADO_CALIDAD': visit_data.get('encargado_calidad', ''),
            'PROFESIONAL_OBRA': visit_data.get('profesional_obra', ''),
            'INSPECTOR_TECNICO': visit_data.get('inspector_tecnico', ''),
            'OTRO_CONTACTO': visit_data.get('otro_contacto', ''),
            
            # Uso de maquinaria
            'MAQUINA_1': visit_data.get('maquinaria', [{}])[0].get('maquina', '') if visit_data.get('maquinaria') else '',
            'INICIO_1': visit_data.get('maquinaria', [{}])[0].get('inicio', '') if visit_data.get('maquinaria') else '',
            'CORTE_1': visit_data.get('maquinaria', [{}])[0].get('corte', '') if visit_data.get('maquinaria') else '',
            'MAQUINA_2': visit_data.get('maquinaria', [{}])[1].get('maquina', '') if len(visit_data.get('maquinaria', [])) > 1 else '',
            'INICIO_2': visit_data.get('maquinaria', [{}])[1].get('inicio', '') if len(visit_data.get('maquinaria', [])) > 1 else '',
            'CORTE_2': visit_data.get('maquinaria', [{}])[1].get('corte', '') if len(visit_data.get('maquinaria', [])) > 1 else '',
            'RETIRO_MAQ': visit_data.get('retiro_maq', 'No'),
            'NUMERO_REPORTE': visit_data.get('numero_reporte', ''),
            
            # Observaciones
            'OBS_MURO_TABIQUE': visit_data.get('obs_muro_tabique', ''),
            'OBS_MATRIZ': visit_data.get('obs_matriz', ''),
            'OBS_LOZA': visit_data.get('obs_loza', ''),
            'OBS_ALMACENAJE': visit_data.get('obs_almacenaje', ''),
            'OBS_PRE_ARMADOS': visit_data.get('obs_pre_armados', ''),
            'OBS_EXTERIORES': visit_data.get('obs_exteriores', ''),
            'OBS_GENERALES': visit_data.get('obs_generales', ''),
            
            # Firmas
            'FIRMA_TECNICO': visit_data.get('firma_tecnico', ''),
            'FIRMA_INSTALADOR': visit_data.get('firma_instalador', ''),
        }
        
        return context
    
    def _generate_polifusion_incident_filename(self, incident_data: Dict) -> str:
        """Generate filename for Polifusión incident documents"""
        now = datetime.now()
        date_str = now.strftime("%Y%m%d")
        
        cliente = incident_data.get('cliente', 'cliente').replace(' ', '_')
        proveedor = incident_data.get('proveedor', 'proveedor').replace(' ', '_')
        
        filename = f"INFORME_INCIDENCIA_POLIFUSION_{cliente}_{proveedor}_{date_str}.docx"
        
        return filename
    
    def _generate_polifusion_visit_filename(self, visit_data: Dict) -> str:
        """Generate filename for Polifusión visit documents"""
        now = datetime.now()
        date_str = now.strftime("%Y%m%d")
        
        obra = visit_data.get('obra', 'obra').replace(' ', '_')
        cliente = visit_data.get('cliente', 'cliente').replace(' ', '_')
        
        filename = f"REPORTE_VISITA_POLIFUSION_{obra}_{cliente}_{date_str}.docx"
        
        return filename

# Global instance
document_generator = DocumentGenerator()
