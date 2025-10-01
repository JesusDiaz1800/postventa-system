"""
Servicio para generar y almacenar documentos en la carpeta compartida
"""
import os
import json
from datetime import datetime
from django.conf import settings
from django.template.loader import render_to_string
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .pdf_generator_service import pdf_generator
import logging

logger = logging.getLogger(__name__)

class DocumentGeneratorService:
    """
    Servicio para generar documentos PDF y DOCX en la carpeta compartida
    """
    
    def __init__(self):
        # Configurar rutas de almacenamiento
        self.shared_folder = getattr(settings, 'SHARED_DOCUMENTS_FOLDER', '/shared/documents')
        self.ensure_directories()
    
    def ensure_directories(self):
        """Crear directorios necesarios si no existen"""
        directories = [
            self.shared_folder,
            os.path.join(self.shared_folder, 'visit_reports'),
            os.path.join(self.shared_folder, 'lab_reports'),
            os.path.join(self.shared_folder, 'supplier_reports'),
            os.path.join(self.shared_folder, 'generated'),
        ]
        
        for directory in directories:
            os.makedirs(directory, exist_ok=True)
    
    def overwrite_existing_document(self, old_filepath, new_content, file_type='pdf'):
        """
        Sobrescribir un documento existente con nuevo contenido
        
        Args:
            old_filepath (str): Ruta del archivo existente
            new_content: Contenido nuevo (bytes para PDF, Document para DOCX)
            file_type (str): Tipo de archivo ('pdf' o 'docx')
        
        Returns:
            str: Ruta del archivo sobrescrito
        """
        try:
            if os.path.exists(old_filepath):
                # Hacer backup del archivo anterior
                backup_path = f"{old_filepath}.backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                os.rename(old_filepath, backup_path)
                logger.info(f"Backup creado: {backup_path}")
            
            # Escribir el nuevo contenido
            if file_type == 'pdf':
                with open(old_filepath, 'wb') as f:
                    f.write(new_content.getvalue())
            elif file_type == 'docx':
                new_content.save(old_filepath)
            
            logger.info(f"Documento sobrescrito exitosamente: {old_filepath}")
            return old_filepath
            
        except Exception as e:
            logger.error(f"Error sobrescribiendo documento: {e}")
            raise
    
    def generate_visit_report_pdf(self, visit_report):
        """Generar documento PDF profesional para reporte de visita"""
        try:
            # Preparar datos del reporte
            report_data = {
                'report_number': visit_report.report_number,
                'project_name': visit_report.project_name,
                'location': visit_report.location,
                'visit_date': visit_report.visit_date.strftime('%d/%m/%Y') if visit_report.visit_date else 'N/A',
                'technician_name': visit_report.technician_name,
                'client_name': visit_report.client_name,
                'client_contact': visit_report.client_contact,
                'objective': visit_report.objective,
                'observations': visit_report.observations,
                'actions_taken': visit_report.actions_taken,
                'recommendations': visit_report.recommendations,
                'conclusions': visit_report.conclusions,
            }
            
            # Preparar datos de la incidencia si existe
            incident_data = None
            if visit_report.incident:
                incident_data = {
                    'code': visit_report.incident.code,
                    'cliente': visit_report.incident.cliente,
                    'obra': visit_report.incident.obra,
                    'provider': visit_report.incident.provider,
                    'categoria': visit_report.incident.categoria,
                    'subcategoria': visit_report.incident.subcategoria,
                    'prioridad': visit_report.incident.prioridad,
                    'estado': visit_report.incident.estado,
                    'fecha_deteccion': visit_report.incident.fecha_deteccion.strftime('%d/%m/%Y') if visit_report.incident.fecha_deteccion else 'N/A',
                    'hora_deteccion': visit_report.incident.hora_deteccion,
                    'sku': visit_report.incident.sku,
                    'lote': visit_report.incident.lote,
                    'factura_num': visit_report.incident.factura_num,
                    'pedido_num': visit_report.incident.pedido_num,
                    'created_by': visit_report.created_by.username if visit_report.created_by else 'Sistema',
                    'document_type': 'Reporte de Visita'
                }
            
            # Generar PDF
            pdf_buffer = pdf_generator.generate_visit_report_pdf(report_data, incident_data)
            
            # Determinar ruta del archivo
            if visit_report.pdf_path and os.path.exists(visit_report.pdf_path):
                # Sobrescribir archivo existente
                filepath = self.overwrite_existing_document(visit_report.pdf_path, pdf_buffer, 'pdf')
                logger.info(f"PDF de reporte de visita sobrescrito: {filepath}")
            else:
                # Crear nuevo archivo
                filename = f"reporte_visita_{visit_report.report_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                filepath = os.path.join(self.shared_folder, 'visit_reports', filename)
                
                with open(filepath, 'wb') as f:
                    f.write(pdf_buffer.getvalue())
                
                logger.info(f"PDF de reporte de visita generado: {filepath}")
            
            # Actualizar modelo
            visit_report.pdf_path = filepath
            visit_report.save(update_fields=['pdf_path'])
            
            return filepath
            
        except Exception as e:
            logger.error(f"Error generando PDF de reporte de visita: {e}")
            raise
    
    def generate_lab_report_pdf(self, lab_report):
        """Generar documento PDF profesional para informe de laboratorio"""
        try:
            # Preparar datos del reporte
            report_data = {
                'report_number': lab_report.report_number,
                'sample_id': lab_report.sample_id,
                'analysis_type': lab_report.analysis_type,
                'received_date': lab_report.received_date.strftime('%d/%m/%Y') if lab_report.received_date else 'N/A',
                'analysis_date': lab_report.analysis_date.strftime('%d/%m/%Y') if lab_report.analysis_date else 'N/A',
                'analyst_name': lab_report.analyst_name,
                'method_used': lab_report.method_used,
                'executive_summary': lab_report.executive_summary,
                'methodology': lab_report.methodology,
                'results': lab_report.results,
                'analysis': lab_report.analysis,
                'conclusions': lab_report.conclusions,
                'recommendations': lab_report.recommendations,
            }
            
            # Preparar datos de la incidencia si existe
            incident_data = None
            if lab_report.incident:
                incident_data = {
                    'code': lab_report.incident.code,
                    'cliente': lab_report.incident.cliente,
                    'obra': lab_report.incident.obra,
                    'provider': lab_report.incident.provider,
                    'categoria': lab_report.incident.categoria,
                    'subcategoria': lab_report.incident.subcategoria,
                    'prioridad': lab_report.incident.prioridad,
                    'estado': lab_report.incident.estado,
                    'fecha_deteccion': lab_report.incident.fecha_deteccion.strftime('%d/%m/%Y') if lab_report.incident.fecha_deteccion else 'N/A',
                    'hora_deteccion': lab_report.incident.hora_deteccion,
                    'sku': lab_report.incident.sku,
                    'lote': lab_report.incident.lote,
                    'factura_num': lab_report.incident.factura_num,
                    'pedido_num': lab_report.incident.pedido_num,
                    'created_by': lab_report.created_by.username if lab_report.created_by else 'Sistema',
                    'document_type': 'Informe de Laboratorio'
                }
            
            # Generar PDF
            pdf_buffer = pdf_generator.generate_lab_report_pdf(report_data, incident_data)
            
            # Determinar ruta del archivo
            if lab_report.pdf_path and os.path.exists(lab_report.pdf_path):
                # Sobrescribir archivo existente
                filepath = self.overwrite_existing_document(lab_report.pdf_path, pdf_buffer, 'pdf')
                logger.info(f"PDF de informe de laboratorio sobrescrito: {filepath}")
            else:
                # Crear nuevo archivo
                filename = f"informe_lab_{lab_report.report_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                filepath = os.path.join(self.shared_folder, 'lab_reports', filename)
                
                with open(filepath, 'wb') as f:
                    f.write(pdf_buffer.getvalue())
                
                logger.info(f"PDF de informe de laboratorio generado: {filepath}")
            
            # Actualizar modelo
            lab_report.pdf_path = filepath
            lab_report.save(update_fields=['pdf_path'])
            
            logger.info(f"PDF de informe de laboratorio generado: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error generando PDF de informe de laboratorio: {e}")
            raise
    
    def generate_supplier_report_pdf(self, supplier_report):
        """Generar documento PDF profesional para informe de proveedor"""
        try:
            # Preparar datos del reporte
            report_data = {
                'report_number': supplier_report.report_number,
                'supplier_name': supplier_report.supplier_name,
                'supplier_contact': supplier_report.supplier_contact,
                'product_name': supplier_report.product_name,
                'batch_number': supplier_report.batch_number,
                'issue_date': supplier_report.issue_date.strftime('%d/%m/%Y') if supplier_report.issue_date else 'N/A',
                'responsible_person': supplier_report.responsible_person,
                'problem_description': supplier_report.problem_description,
                'analysis_performed': supplier_report.analysis_performed,
                'findings': supplier_report.findings,
                'required_actions': supplier_report.required_actions,
                'deadlines': supplier_report.deadlines,
                'contact_info': supplier_report.contact_info,
            }
            
            # Preparar datos de la incidencia si existe
            incident_data = None
            if supplier_report.incident:
                incident_data = {
                    'code': supplier_report.incident.code,
                    'cliente': supplier_report.incident.cliente,
                    'obra': supplier_report.incident.obra,
                    'provider': supplier_report.incident.provider,
                    'categoria': supplier_report.incident.categoria,
                    'subcategoria': supplier_report.incident.subcategoria,
                    'prioridad': supplier_report.incident.prioridad,
                    'estado': supplier_report.incident.estado,
                    'fecha_deteccion': supplier_report.incident.fecha_deteccion.strftime('%d/%m/%Y') if supplier_report.incident.fecha_deteccion else 'N/A',
                    'hora_deteccion': supplier_report.incident.hora_deteccion,
                    'sku': supplier_report.incident.sku,
                    'lote': supplier_report.incident.lote,
                    'factura_num': supplier_report.incident.factura_num,
                    'pedido_num': supplier_report.incident.pedido_num,
                    'created_by': supplier_report.created_by.username if supplier_report.created_by else 'Sistema',
                    'document_type': 'Informe para Proveedor'
                }
            
            # Generar PDF
            pdf_buffer = pdf_generator.generate_supplier_report_pdf(report_data, incident_data)
            
            # Determinar ruta del archivo
            if supplier_report.pdf_path and os.path.exists(supplier_report.pdf_path):
                # Sobrescribir archivo existente
                filepath = self.overwrite_existing_document(supplier_report.pdf_path, pdf_buffer, 'pdf')
                logger.info(f"PDF de informe de proveedor sobrescrito: {filepath}")
            else:
                # Crear nuevo archivo
                filename = f"informe_proveedor_{supplier_report.report_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                filepath = os.path.join(self.shared_folder, 'supplier_reports', filename)
                
                with open(filepath, 'wb') as f:
                    f.write(pdf_buffer.getvalue())
                
                logger.info(f"PDF de informe de proveedor generado: {filepath}")
            
            # Actualizar modelo
            supplier_report.pdf_path = filepath
            supplier_report.save(update_fields=['pdf_path'])
            
            return filepath
            
        except Exception as e:
            logger.error(f"Error generando PDF de informe de proveedor: {e}")
            raise

    def generate_visit_report_docx(self, visit_report):
        """Generar documento DOCX para reporte de visita"""
        try:
            # Crear documento
            doc = Document()
            
            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Título
            title = doc.add_heading('REPORTE VISITA', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información de la empresa
            company_info = doc.add_paragraph()
            company_info.add_run('POLIFUSIÓN S.A.\n').bold = True
            company_info.add_run('Especialistas en Tuberías y Fittings de Polipropileno\n')
            company_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Fecha
            date_para = doc.add_paragraph()
            date_para.add_run(f'Fecha: {visit_report.visit_date.strftime("%d/%m/%Y")}\n')
            date_para.add_run(f'Orden N°: {visit_report.order_number}\n')
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Información del proyecto
            doc.add_heading('INFORMACIÓN DEL PROYECTO', level=1)
            
            project_table = doc.add_table(rows=8, cols=2)
            project_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            project_table.style = 'Table Grid'
            
            # Configurar encabezados
            hdr_cells = project_table.rows[0].cells
            hdr_cells[0].text = 'Campo'
            hdr_cells[1].text = 'Valor'
            
            # Datos del proyecto
            project_data = [
                ('Obra', visit_report.project_name),
                ('ID Proyecto', visit_report.project_id or 'N/A'),
                ('Cliente', visit_report.client_name),
                ('RUT Cliente', visit_report.client_rut or 'N/A'),
                ('Dirección', visit_report.address),
                ('Empresa Constructora', visit_report.construction_company or 'N/A'),
                ('Motivo Visita', visit_report.visit_reason),
            ]
            
            for i, (field, value) in enumerate(project_data, 1):
                row = project_table.rows[i]
                row.cells[0].text = field
                row.cells[1].text = str(value)
            
            # Personal involucrado
            doc.add_heading('PERSONAL INVOLUCRADO', level=1)
            
            personnel_table = doc.add_table(rows=5, cols=2)
            personnel_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            personnel_table.style = 'Table Grid'
            
            hdr_cells = personnel_table.rows[0].cells
            hdr_cells[0].text = 'Rol'
            hdr_cells[1].text = 'Nombre'
            
            personnel_data = [
                ('Vendedor', visit_report.salesperson),
                ('Técnico', visit_report.technician),
                ('Instalador', visit_report.installer or 'N/A'),
                ('Teléfono Instalador', visit_report.installer_phone or 'N/A'),
            ]
            
            for i, (role, name) in enumerate(personnel_data, 1):
                row = personnel_table.rows[i]
                row.cells[0].text = role
                row.cells[1].text = str(name)
            
            # Ubicación
            doc.add_heading('UBICACIÓN', level=1)
            location_para = doc.add_paragraph()
            location_para.add_run(f'Comuna: {visit_report.commune}\n')
            location_para.add_run(f'Ciudad: {visit_report.city}\n')
            
            # Observaciones
            if visit_report.general_observations:
                doc.add_heading('OBSERVACIONES GENERALES', level=1)
                obs_para = doc.add_paragraph(visit_report.general_observations)
            
            # Observaciones por sección
            sections_obs = [
                ('OBS Muro/Tabique', visit_report.wall_observations),
                ('OBS Matriz', visit_report.matrix_observations),
                ('OBS Losa', visit_report.slab_observations),
                ('OBS Almacenaje', visit_report.storage_observations),
                ('OBS Pre Armados', visit_report.pre_assembled_observations),
                ('OBS Exteriores', visit_report.exterior_observations),
            ]
            
            for section_title, section_obs in sections_obs:
                if section_obs:
                    doc.add_heading(section_title, level=2)
                    doc.add_paragraph(section_obs)
            
            # Firmas
            doc.add_heading('FIRMAS', level=1)
            signatures_table = doc.add_table(rows=3, cols=2)
            signatures_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            signatures_table.style = 'Table Grid'
            
            hdr_cells = signatures_table.rows[0].cells
            hdr_cells[0].text = 'Firma Técnico'
            hdr_cells[1].text = 'Firma Instalador'
            
            signatures_table.rows[1].cells[0].text = '_________________'
            signatures_table.rows[1].cells[1].text = '_________________'
            signatures_table.rows[2].cells[0].text = visit_report.technician
            signatures_table.rows[2].cells[1].text = visit_report.installer or 'N/A'
            
            # Guardar documento
            filename = f"RV_{visit_report.report_number}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath = os.path.join(self.shared_folder, 'visit_reports', filename)
            doc.save(filepath)
            
            logger.info(f"Reporte de visita generado: {filepath}")
            return filepath, filename
            
        except Exception as e:
            logger.error(f"Error generando reporte de visita: {e}")
            raise
    
    def generate_lab_report_docx(self, lab_report):
        """Generar documento DOCX para informe de laboratorio"""
        try:
            # Crear documento
            doc = Document()
            
            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Header con logo
            header_para = doc.add_paragraph()
            header_para.add_run('POLIFUSIÓN S.A.\n').bold = True
            header_para.add_run('Especialistas en Tuberías y Fittings de Polipropileno\n')
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Título
            title = doc.add_heading('Informe de laboratorio', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del formulario
            form_info = doc.add_paragraph()
            form_info.add_run(f'Departamento de control de calidad\n')
            form_info.add_run(f'{lab_report.form_number}\n')
            form_info.add_run(f'En uso desde el: 28-08-2013\n')
            form_info.add_run(f'Página 1 de 1\n')
            form_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información general
            doc.add_heading('INFORMACIÓN GENERAL', level=1)
            
            general_table = doc.add_table(rows=4, cols=2)
            general_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            general_table.style = 'Table Grid'
            
            general_data = [
                ('SOLICITANTE', lab_report.applicant),
                ('CLIENTE', lab_report.client),
                ('Fecha de solicitud', lab_report.request_date.strftime('%d/%m/%Y')),
                ('Incidencia relacionada', lab_report.related_incident.code if lab_report.related_incident else 'N/A'),
            ]
            
            for i, (field, value) in enumerate(general_data):
                row = general_table.rows[i]
                row.cells[0].text = field
                row.cells[1].text = str(value)
            
            # Descripción
            doc.add_heading('DESCRIPCIÓN', level=1)
            desc_para = doc.add_paragraph(lab_report.description)
            
            # Antecedentes del proyecto
            if lab_report.project_background:
                doc.add_heading('ANTECEDENTES DEL PROYECTO', level=1)
                bg_para = doc.add_paragraph(lab_report.project_background)
            
            # Ensayos realizados
            if lab_report.tests_performed:
                doc.add_heading('ENSAYOS REALIZADOS', level=1)
                for test in lab_report.tests_performed:
                    if test:
                        doc.add_paragraph(f'• {test}')
            
            # Comentarios
            if lab_report.comments:
                doc.add_heading('COMENTARIOS', level=1)
                comments_para = doc.add_paragraph(lab_report.comments)
            
            # Conclusiones
            if lab_report.conclusions:
                doc.add_heading('CONCLUSIONES', level=1)
                conclusions_para = doc.add_paragraph(lab_report.conclusions)
            
            # Recomendaciones
            if lab_report.recommendations:
                doc.add_heading('RECOMENDACIONES', level=1)
                recommendations_para = doc.add_paragraph(lab_report.recommendations)
            
            # Firma del experto técnico
            if lab_report.technical_expert_name:
                doc.add_heading('FIRMA DEL EXPERTO TÉCNICO', level=1)
                expert_para = doc.add_paragraph()
                expert_para.add_run('_________________\n')
                expert_para.add_run(lab_report.technical_expert_name)
                expert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Guardar documento
            filename = f"IL_{lab_report.report_number}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath = os.path.join(self.shared_folder, 'lab_reports', filename)
            doc.save(filepath)
            
            logger.info(f"Informe de laboratorio generado: {filepath}")
            return filepath, filename
            
        except Exception as e:
            logger.error(f"Error generando informe de laboratorio: {e}")
            raise
    
    def generate_supplier_report_docx(self, supplier_report):
        """Generar documento DOCX para informe de proveedor"""
        try:
            # Crear documento
            doc = Document()
            
            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Header
            header_para = doc.add_paragraph()
            header_para.add_run('POLIFUSIÓN S.A.\n').bold = True
            header_para.add_run('Especialistas en Tuberías y Fittings de Polipropileno\n')
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Fecha
            date_para = doc.add_paragraph()
            date_para.add_run(f'Fecha: {supplier_report.report_date.strftime("%d/%m/%Y")}\n')
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Información del proveedor
            doc.add_heading('INFORMACIÓN DEL PROVEEDOR', level=1)
            
            supplier_table = doc.add_table(rows=4, cols=2)
            supplier_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            supplier_table.style = 'Table Grid'
            
            supplier_data = [
                ('Proveedor', supplier_report.supplier_name),
                ('Contacto', supplier_report.supplier_contact or 'N/A'),
                ('Email', supplier_report.supplier_email or 'N/A'),
                ('Incidencia relacionada', supplier_report.related_incident.code if supplier_report.related_incident else 'N/A'),
            ]
            
            for i, (field, value) in enumerate(supplier_data):
                row = supplier_table.rows[i]
                row.cells[0].text = field
                row.cells[1].text = str(value)
            
            # Asunto
            doc.add_heading('ASUNTO', level=1)
            subject_para = doc.add_paragraph(supplier_report.subject)
            
            # Introducción
            doc.add_heading('INTRODUCCIÓN', level=1)
            intro_para = doc.add_paragraph(supplier_report.introduction)
            
            # Descripción del problema
            doc.add_heading('DESCRIPCIÓN DEL PROBLEMA', level=1)
            problem_para = doc.add_paragraph(supplier_report.problem_description)
            
            # Análisis técnico
            doc.add_heading('ANÁLISIS TÉCNICO', level=1)
            analysis_para = doc.add_paragraph(supplier_report.technical_analysis)
            
            # Recomendaciones
            doc.add_heading('RECOMENDACIONES', level=1)
            recommendations_para = doc.add_paragraph(supplier_report.recommendations)
            
            # Mejoras esperadas
            doc.add_heading('MEJORAS ESPERADAS', level=1)
            improvements_para = doc.add_paragraph(supplier_report.expected_improvements)
            
            # Firma
            doc.add_heading('FIRMA', level=1)
            signature_para = doc.add_paragraph()
            signature_para.add_run('_________________\n')
            signature_para.add_run('POLIFUSIÓN S.A.')
            signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Guardar documento
            filename = f"IP_{supplier_report.report_number}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath = os.path.join(self.shared_folder, 'supplier_reports', filename)
            doc.save(filepath)
            
            logger.info(f"Informe para proveedor generado: {filepath}")
            return filepath, filename
            
        except Exception as e:
            logger.error(f"Error generando informe para proveedor: {e}")
            raise
    
    def get_document_path(self, document_type, filename):
        """Obtener la ruta completa de un documento"""
        return os.path.join(self.shared_folder, document_type, filename)
    
    def list_documents(self, document_type=None):
        """Listar documentos en la carpeta compartida"""
        documents = []
        
        if document_type:
            types = [document_type]
        else:
            types = ['visit_reports', 'lab_reports', 'supplier_reports']
        
        for doc_type in types:
            folder_path = os.path.join(self.shared_folder, doc_type)
            if os.path.exists(folder_path):
                for filename in os.listdir(folder_path):
                    if filename.endswith(('.docx', '.pdf')):
                        filepath = os.path.join(folder_path, filename)
                        stat = os.stat(filepath)
                        documents.append({
                            'type': doc_type,
                            'filename': filename,
                            'filepath': filepath,
                            'size': stat.st_size,
                            'created': datetime.fromtimestamp(stat.st_ctime),
                            'modified': datetime.fromtimestamp(stat.st_mtime),
                        })
        
        return sorted(documents, key=lambda x: x['modified'], reverse=True)
    
    def delete_document(self, document_type, filename):
        """Eliminar un documento de la carpeta compartida"""
        filepath = os.path.join(self.shared_folder, document_type, filename)
        if os.path.exists(filepath):
            os.remove(filepath)
            logger.info(f"Documento eliminado: {filepath}")
            return True
        return False
