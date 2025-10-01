"""
Plantillas específicas de Polifusión basadas en los informes de laboratorio
Genera documentos con el formato exacto de la empresa
"""

import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class PolifusionTemplateGenerator:
    """Generador de plantillas específicas de Polifusión"""
    
    def __init__(self):
        self.company_name = "Polifusión"
        self.company_tagline = "Especialistas en Tuberías y Fittings de Polipropileno"
        self.department = "Departamento de control de Calidad"
        self.form_id = "FORM-SERTEC-0007"
        self.form_date = "En uso desde el: 28-08-2013"
    
    def create_lab_report_template(self, output_path: str):
        """Crear plantilla de informe de laboratorio de Polifusión"""
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header con logo y información de la empresa
        self._add_header(doc)
        
        # Información del solicitante
        self._add_applicant_info(doc)
        
        # Descripción
        self._add_description_section(doc)
        
        # Ensayos realizados
        self._add_tests_section(doc)
        
        # Comentarios
        self._add_comments_section(doc)
        
        # Conclusiones
        self._add_conclusions_section(doc)
        
        # Guardar plantilla
        doc.save(output_path)
        return output_path
    
    def _add_header(self, doc):
        """Agregar header con logo y información de la empresa"""
        # Tabla para el header
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Configurar columnas
        cols = table.columns
        cols[0].width = Inches(2.5)
        cols[1].width = Inches(3)
        cols[2].width = Inches(2.5)
        
        # Celda izquierda - Logo y nombre de empresa
        left_cell = table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Logo placeholder (se reemplazará con imagen real)
        logo_para = left_cell.add_paragraph()
        logo_para.add_run("LOGO POLIFUSIÓN").font.size = Pt(14)
        logo_para.add_run("\n").font.size = Pt(8)
        logo_para.add_run(self.company_tagline).font.size = Pt(10)
        
        # Celda central - Título del documento
        center_cell = table.cell(0, 1)
        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        dept_para = center_cell.add_paragraph()
        dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dept_para.add_run(self.department).font.size = Pt(12)
        dept_para.add_run("\n").font.size = Pt(8)
        
        title_para = center_cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.add_run("Informe de laboratorio").font.size = Pt(14)
        title_para.add_run("\n").font.size = Pt(8)
        
        # Celda derecha - Información del formulario
        right_cell = table.cell(0, 2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        form_para = right_cell.add_paragraph()
        form_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        form_para.add_run(self.form_id).font.size = Pt(10)
        form_para.add_run("\n").font.size = Pt(8)
        form_para.add_run(self.form_date).font.size = Pt(8)
        form_para.add_run("\n").font.size = Pt(8)
        form_para.add_run("Página 1 de 1").font.size = Pt(8)
        
        # Espacio después del header
        doc.add_paragraph()
    
    def _add_applicant_info(self, doc):
        """Agregar información del solicitante"""
        # Título de sección
        title = doc.add_heading('INFORMACIÓN DEL SOLICITANTE', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tabla de información
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # Configurar columnas
        cols = table.columns
        cols[0].width = Inches(2)
        cols[1].width = Inches(4)
        
        # Datos del solicitante
        data = [
            ("SOLICITANTE:", "{{SOLICITANTE}}"),
            ("Fecha de solicitud:", "{{FECHA_SOLICITUD}}"),
            ("CLIENTE:", "{{CLIENTE}}")
        ]
        
        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph()
    
    def _add_description_section(self, doc):
        """Agregar sección de descripción"""
        title = doc.add_heading('DESCRIPCIÓN', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Contenido de descripción con placeholders
        desc_para = doc.add_paragraph()
        desc_para.add_run("Se recibió una muestra compuesta por un codo de 90° PP-RCT de {{DIAMETRO}}mm, fusionado en ambos extremos a una tubería PP-RCT/FIBERGLASS S-3.2 de {{DIAMETRO}}mm.\n\n")
        desc_para.add_run("La muestra presenta una grieta longitudinal tipo corte en la sección transversal de la tubería en uno de los extremos del codo, con un inserto de tapón, posiblemente para detener una fuga.\n\n")
        desc_para.add_run("Se realizó una caracterización detallada, seguida de un corte transversal para evaluar la zona de ruptura y la calidad de la fusión.\n\n")
        desc_para.add_run("Proyecto: {{PROYECTO}}\n")
        desc_para.add_run("Ubicación: {{UBICACION}}\n")
        desc_para.add_run("Presión registrada: {{PRESION}} bar\n")
        desc_para.add_run("Temperatura: {{TEMPERATURA}}\n")
        desc_para.add_run("Información proporcionada por: {{INFORMANTE}}")
        
        doc.add_paragraph()
    
    def _add_tests_section(self, doc):
        """Agregar sección de ensayos realizados"""
        title = doc.add_heading('ENSAYOS REALIZADOS', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        tests_para = doc.add_paragraph()
        tests_para.add_run("• Evaluación visual de la muestra proporcionada.\n")
        tests_para.add_run("• Peritaje de memoria térmica del cordón de fusión.\n")
        tests_para.add_run("• Análisis de la zona de ruptura.\n")
        tests_para.add_run("• Evaluación de la calidad de la fusión.\n")
        tests_para.add_run("• {{ENSAYOS_ADICIONALES}}")
        
        doc.add_paragraph()
    
    def _add_comments_section(self, doc):
        """Agregar sección de comentarios"""
        title = doc.add_heading('COMENTARIOS', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        comments_para = doc.add_paragraph()
        comments_para.add_run("{{COMENTARIOS_DETALLADOS}}\n\n")
        comments_para.add_run("Se observó evidencia de presencia de agua durante el proceso de fusión, indicado por cordones de fusión internos corrugados.\n\n")
        comments_para.add_run("La inspección de la grieta reveló una separación de 1-2 mm de la pared del fitting, justo detrás del cordón de fusión.\n\n")
        comments_para.add_run("Al aplicar una fuerza de flexión se reveló una fractura frágil característica, típica de material cristalizado debido al calentamiento excesivo.")
        
        doc.add_paragraph()
    
    def _add_conclusions_section(self, doc):
        """Agregar sección de conclusiones"""
        title = doc.add_heading('CONCLUSIONES', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        conclusions_para = doc.add_paragraph()
        conclusions_para.add_run("{{CONCLUSIONES_DETALLADAS}}\n\n")
        conclusions_para.add_run("Basado en la evidencia observada, el proceso de ensamble y termofusión fue inadecuado.\n\n")
        conclusions_para.add_run("Los tiempos de fusión excesivos contribuyeron a la cristalización del material y tiempos mínimos de fusión. La presencia de agua aceleró el proceso de enfriamiento, creando tensión entre los materiales fusionados, lo que finalmente llevó a la formación de una grieta con fractura frágil característica en el material cristalizado.")
        
        # Firma
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tabla para firma
        sig_table = doc.add_table(rows=1, cols=1)
        sig_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        sig_cell = sig_table.cell(0, 0)
        sig_para = sig_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        sig_para.add_run("_________________________\n")
        sig_para.add_run("{{EXPERTO_NOMBRE}}\n")
        sig_para.add_run("Experto técnico\n")
        sig_para.add_run("Control de Calidad Polifusión S.A.")
    
    def create_photographs_template(self, output_path: str):
        """Crear plantilla para página de fotografías"""
        doc = Document()
        
        # Header
        self._add_header(doc)
        
        # Título de fotografías
        title = doc.add_heading('FOTOGRAFÍAS', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Sección de fotografías
        photos_para = doc.add_paragraph()
        photos_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        photos_para.add_run("MUESTRA RECIBIDA").font.size = Pt(12)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Placeholder para imagen 1
        img1_para = doc.add_paragraph()
        img1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img1_para.add_run("[FOTOGRAFÍA 1: MUESTRA RECIBIDA]").font.size = Pt(10)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Descripción de la segunda foto
        desc2_para = doc.add_paragraph()
        desc2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc2_para.add_run("RUPTURA LONGITUDINAL TIPO GRIETA, LA CUAL FUE INTERVENIDA CON TARUGO DE REPARACIÓN").font.size = Pt(12)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Placeholder para imagen 2
        img2_para = doc.add_paragraph()
        img2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img2_para.add_run("[FOTOGRAFÍA 2: RUPTURA Y REPARACIÓN]").font.size = Pt(10)
        
        doc.save(output_path)
        return output_path
    
    def create_detailed_analysis_template(self, output_path: str):
        """Crear plantilla para análisis detallado"""
        doc = Document()
        
        # Header
        self._add_header(doc)
        
        # Contenido del análisis detallado
        content_para = doc.add_paragraph()
        content_para.add_run("{{ANALISIS_DETALLADO}}\n\n")
        content_para.add_run("CORDÓN DE SOLDADURA EXTERIOR IRREGULAR\n\n")
        content_para.add_run("SOLDADURA EXTERIOR IRREGULAR, INDICA SOBRE CALENTAMIENTO EXCESIVO DE TIEMPO Y TEMPERATURA.")
        
        # Placeholders para imágenes
        doc.add_paragraph()
        img1_para = doc.add_paragraph()
        img1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img1_para.add_run("[FOTOGRAFÍA: CORDÓN DE SOLDADURA IRREGULAR]").font.size = Pt(10)
        
        doc.add_paragraph()
        img2_para = doc.add_paragraph()
        img2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img2_para.add_run("[FOTOGRAFÍA: SOBRECALENTAMIENTO]").font.size = Pt(10)
        
        doc.save(output_path)
        return output_path
    
    def create_incident_report_template(self, output_path: str):
        """Crear plantilla de informe de incidencia de Polifusión"""
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header con logo y información de la empresa
        self._add_incident_header(doc)
        
        # Registro de información
        self._add_incident_info(doc)
        
        # Descripción del problema
        self._add_problem_description(doc)
        
        # Acciones inmediatas
        self._add_immediate_actions(doc)
        
        # Evolución/Acciones posteriores
        self._add_evolution_actions(doc)
        
        # Observaciones y cierre
        self._add_observations_closure(doc)
        
        # Guardar plantilla
        doc.save(output_path)
        return output_path
    
    def create_visit_report_template(self, output_path: str):
        """Crear plantilla de reporte de visita de Polifusión"""
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header con logo y información de la empresa
        self._add_visit_header(doc)
        
        # Información general de la obra
        self._add_work_info(doc)
        
        # Información del personal
        self._add_personnel_info(doc)
        
        # Roles y contactos
        self._add_roles_contacts(doc)
        
        # Uso de maquinaria
        self._add_machinery_usage(doc)
        
        # Observaciones
        self._add_observations(doc)
        
        # Firmas
        self._add_signatures(doc)
        
        # Guardar plantilla
        doc.save(output_path)
        return output_path
    
    def _add_incident_header(self, doc):
        """Agregar header para informe de incidencia"""
        # Tabla para el header
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Configurar columnas
        cols = table.columns
        cols[0].width = Inches(2.5)
        cols[1].width = Inches(3)
        cols[2].width = Inches(2.5)
        
        # Celda izquierda - Logo y nombre de empresa
        left_cell = table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Logo placeholder
        logo_para = left_cell.add_paragraph()
        logo_para.add_run("LOGO POLIFUSIÓN").font.size = Pt(14)
        logo_para.add_run("\n").font.size = Pt(8)
        logo_para.add_run(self.company_tagline).font.size = Pt(10)
        
        # Celda central - Título del documento
        center_cell = table.cell(0, 1)
        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_para = center_cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.add_run("INFORME DE INCIDENCIA").font.size = Pt(16)
        
        # Celda derecha - Información del formulario
        right_cell = table.cell(0, 2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        form_para = right_cell.add_paragraph()
        form_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        form_para.add_run("INFORMES 800").font.size = Pt(12)
        form_para.add_run("\n").font.size = Pt(8)
        form_para.add_run("N° {{REPORT_NUMBER}}").font.size = Pt(10)
        
        doc.add_paragraph()
    
    def _add_incident_info(self, doc):
        """Agregar información del incidente"""
        title = doc.add_heading('REGISTRO DE INFORMACIÓN', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tabla de información
        table = doc.add_table(rows=11, cols=2)
        table.style = 'Table Grid'
        
        # Configurar columnas
        cols = table.columns
        cols[0].width = Inches(2)
        cols[1].width = Inches(4)
        
        # Datos del incidente
        data = [
            ("Proveedor:", "{{PROVEEDOR}}"),
            ("Obra:", "{{OBRA}}"),
            ("Producción:", "{{PRODUCCION}}"),
            ("Cliente:", "{{CLIENTE}}"),
            ("Servicio:", "{{SERVICIO}}"),
            ("RUT:", "{{RUT}}"),
            ("Dirección:", "{{DIRECCION}}"),
            ("Otros:", "{{OTROS}}"),
            ("Contactos:", "{{CONTACTOS}}"),
            ("FECHA DETECCIÓN:", "{{FECHA_DETECCION}}"),
            ("HORA:", "{{HORA}}")
        ]
        
        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph()
    
    def _add_problem_description(self, doc):
        """Agregar descripción del problema"""
        title = doc.add_heading('DESCRIPCIÓN DEL PROBLEMA', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        desc_para = doc.add_paragraph()
        desc_para.add_run("{{DESCRIPCION_PROBLEMA}}")
        
        doc.add_paragraph()
    
    def _add_immediate_actions(self, doc):
        """Agregar acciones inmediatas"""
        title = doc.add_heading('ACCIONES INMEDIATAS ADOPTADAS', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        actions_para = doc.add_paragraph()
        actions_para.add_run("{{ACCIONES_INMEDIATAS}}")
        
        doc.add_paragraph()
    
    def _add_evolution_actions(self, doc):
        """Agregar evolución/acciones posteriores"""
        title = doc.add_heading('DESCRIPCIÓN DE LA EVOLUCIÓN/ACCIONES POSTERIORMENTE ADOPTADAS', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tabla para evolución
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Encabezados
        table.cell(0, 0).text = "FECHA"
        table.cell(0, 1).text = "DESCRIPCIÓN DE LA EVOLUCIÓN/ACCIONES POSTERIORMENTE ADOPTADAS"
        
        # Agregar filas de ejemplo
        for i in range(3):
            row = table.add_row()
            row.cells[0].text = "{{FECHA_" + str(i+1) + "}}"
            row.cells[1].text = "{{DESCRIPCION_" + str(i+1) + "}}"
        
        doc.add_paragraph()
    
    def _add_observations_closure(self, doc):
        """Agregar observaciones y cierre"""
        title = doc.add_heading('OBSERVACIONES Y CIERRE', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        obs_para = doc.add_paragraph()
        obs_para.add_run("OBS: {{OBSERVACIONES}}")
        
        doc.add_paragraph()
        
        fecha_para = doc.add_paragraph()
        fecha_para.add_run("FECHA CIERRE: {{FECHA_CIERRE}}")
        
        doc.add_paragraph()
    
    def _add_visit_header(self, doc):
        """Agregar header para reporte de visita"""
        # Tabla para el header
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Configurar columnas
        cols = table.columns
        cols[0].width = Inches(2.5)
        cols[1].width = Inches(3)
        cols[2].width = Inches(2.5)
        
        # Celda izquierda - Logo y nombre de empresa
        left_cell = table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Logo placeholder
        logo_para = left_cell.add_paragraph()
        logo_para.add_run("LOGO POLIFUSIÓN").font.size = Pt(14)
        logo_para.add_run("\n").font.size = Pt(8)
        logo_para.add_run("POLIFUSION S.A.").font.size = Pt(10)
        logo_para.add_run("\n").font.size = Pt(8)
        logo_para.add_run("Cacique Colin 2525 Lampa").font.size = Pt(9)
        
        # Celda central - Título del documento
        center_cell = table.cell(0, 1)
        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_para = center_cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.add_run("REPORTE VISITA").font.size = Pt(16)
        
        # Celda derecha - Información del formulario
        right_cell = table.cell(0, 2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        form_para = right_cell.add_paragraph()
        form_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        form_para.add_run("Orden N°: {{ORDEN_NUMBER}}").font.size = Pt(10)
        form_para.add_run("\n").font.size = Pt(8)
        form_para.add_run("Fecha Visita: {{FECHA_VISITA}}").font.size = Pt(10)
        
        doc.add_paragraph()
    
    def _add_work_info(self, doc):
        """Agregar información de la obra"""
        title = doc.add_heading('INFORMACIÓN GENERAL DE LA OBRA', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tabla de información
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Configurar columnas
        cols = table.columns
        cols[0].width = Inches(2)
        cols[1].width = Inches(4)
        
        # Datos de la obra
        data = [
            ("Obra:", "{{OBRA}}"),
            ("Cliente:", "{{CLIENTE}}"),
            ("Dir:", "{{DIRECCION}}"),
            ("Adm:", "{{ADMINISTRADOR}}"),
            ("Const:", "{{CONSTRUCTOR}}"),
            ("Motivo V:", "{{MOTIVO_VISITA}}")
        ]
        
        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph()
    
    def _add_personnel_info(self, doc):
        """Agregar información del personal"""
        title = doc.add_heading('INFORMACIÓN DEL PERSONAL', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tabla de información
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Configurar columnas
        cols = table.columns
        cols[0].width = Inches(2)
        cols[1].width = Inches(4)
        
        # Datos del personal
        data = [
            ("Vendedor:", "{{VENDEDOR}}"),
            ("Comuna:", "{{COMUNA}}"),
            ("Ciudad:", "{{CIUDAD}}"),
            ("Instalador:", "{{INSTALADOR}}"),
            ("Fono Inst.:", "{{FONO_INSTALADOR}}"),
            ("Técnico:", "{{TECNICO}}")
        ]
        
        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph()
    
    def _add_roles_contacts(self, doc):
        """Agregar roles y contactos"""
        title = doc.add_heading('ROLES Y CONTACTOS', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tabla de roles
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Encabezados
        table.cell(0, 0).text = "Enc. de Calidad"
        table.cell(0, 1).text = "Profesional de Obra"
        table.cell(0, 2).text = "Inspector Técnico de Obra"
        table.cell(0, 3).text = "Otro Contacto"
        
        # Agregar fila de datos
        row = table.add_row()
        row.cells[0].text = "{{ENCARGADO_CALIDAD}}"
        row.cells[1].text = "{{PROFESIONAL_OBRA}}"
        row.cells[2].text = "{{INSPECTOR_TECNICO}}"
        row.cells[3].text = "{{OTRO_CONTACTO}}"
        
        doc.add_paragraph()
    
    def _add_machinery_usage(self, doc):
        """Agregar uso de maquinaria"""
        title = doc.add_heading('USO DE MAQUINARIA', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tabla de maquinaria
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Encabezados
        table.cell(0, 0).text = "Máquina"
        table.cell(0, 1).text = "Inicio"
        table.cell(0, 2).text = "Corte"
        
        # Agregar filas de ejemplo
        for i in range(2):
            row = table.add_row()
            row.cells[0].text = "{{MAQUINA_" + str(i+1) + "}}"
            row.cells[1].text = "{{INICIO_" + str(i+1) + "}}"
            row.cells[2].text = "{{CORTE_" + str(i+1) + "}}"
        
        doc.add_paragraph()
        
        # Información adicional
        retiro_para = doc.add_paragraph()
        retiro_para.add_run("Retiro Maq: {{RETIRO_MAQ}}")
        
        num_para = doc.add_paragraph()
        num_para.add_run("N° Rep: {{NUMERO_REPORTE}}")
        
        doc.add_paragraph()
    
    def _add_observations(self, doc):
        """Agregar observaciones"""
        title = doc.add_heading('OBSERVACIONES', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Observaciones por categoría
        obs_categories = [
            ("Obs Muro/Tabique:", "{{OBS_MURO_TABIQUE}}"),
            ("Obs Matriz:", "{{OBS_MATRIZ}}"),
            ("Obs Loza:", "{{OBS_LOZA}}"),
            ("Obs Almacenaje:", "{{OBS_ALMACENAJE}}"),
            ("Obs Pre Armados:", "{{OBS_PRE_ARMADOS}}"),
            ("Obs Exteriores:", "{{OBS_EXTERIORES}}"),
            ("Obs Generales:", "{{OBS_GENERALES}}")
        ]
        
        for category, placeholder in obs_categories:
            cat_title = doc.add_heading(category, level=3)
            cat_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            obs_para = doc.add_paragraph()
            obs_para.add_run(placeholder)
            
            doc.add_paragraph()
    
    def _add_signatures(self, doc):
        """Agregar firmas"""
        title = doc.add_heading('FIRMAS', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tabla para firmas
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Encabezados
        table.cell(0, 0).text = "Firma Técnico"
        table.cell(0, 1).text = "Firma Instalador"
        
        # Agregar fila de datos
        row = table.add_row()
        row.cells[0].text = "{{FIRMA_TECNICO}}"
        row.cells[1].text = "{{FIRMA_INSTALADOR}}"
        
        doc.add_paragraph()

# Instancia global
polifusion_generator = PolifusionTemplateGenerator()
