"""
Plantillas mejoradas de Polifusión con diseño profesional y logo integrado
Genera documentos con el formato exacto de la empresa y diseño moderno
"""

import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class PolifusionEnhancedTemplateGenerator:
    """Generador de plantillas mejoradas de Polifusión con diseño profesional"""
    
    def __init__(self):
        self.company_name = "POLIFUSIÓN"
        self.company_tagline = "Especialistas en Tuberías y Fittings de Polipropileno"
        self.department = "Departamento de Control de Calidad"
        self.form_id = "FORM-SERTEC-0007"
        self.form_date = "En uso desde: 28-08-2013"
        
        # Colores corporativos
        self.primary_color = RGBColor(0, 51, 102)      # Azul oscuro
        self.secondary_color = RGBColor(0, 102, 204)   # Azul medio
        self.accent_color = RGBColor(255, 140, 0)      # Naranja
        self.text_color = RGBColor(64, 64, 64)         # Gris oscuro
        self.light_gray = RGBColor(128, 128, 128)      # Gris claro
    
    def create_lab_report_template(self, output_path: str):
        """Crear plantilla mejorada de informe de laboratorio"""
        doc = Document()
        
        # Configurar márgenes profesionales
        self._setup_document_margins(doc)
        
        # Header profesional con logo
        self._add_professional_header(doc)
        
        # Título principal
        self._add_main_title(doc, "INFORME DE LABORATORIO")
        
        # Subtítulo
        self._add_subtitle(doc, "ANÁLISIS TÉCNICO DE TUBERÍAS PP-RCT")
        
        # Información del solicitante
        self._add_applicant_info_section(doc)
        
        # Descripción de la muestra
        self._add_sample_description_section(doc)
        
        # Ensayos realizados
        self._add_tests_section(doc)
        
        # Comentarios detallados
        self._add_detailed_comments_section(doc)
        
        # Conclusiones
        self._add_conclusions_section(doc)
        
        # Firma del experto
        self._add_expert_signature_section(doc)
        
        # Footer profesional
        self._add_professional_footer(doc)
        
        doc.save(output_path)
        return output_path
    
    def create_incident_report_template(self, output_path: str):
        """Crear plantilla mejorada de informe de incidencia"""
        doc = Document()
        
        self._setup_document_margins(doc)
        self._add_professional_header(doc)
        
        self._add_main_title(doc, "INFORME DE INCIDENCIA")
        self._add_subtitle(doc, "REPORTE TÉCNICO DE FALLA EN TUBERÍAS")
        
        # Información de la incidencia
        self._add_incident_info_section(doc)
        
        # Descripción del problema
        self._add_problem_description_section(doc)
        
        # Acciones inmediatas
        self._add_immediate_actions_section(doc)
        
        # Evolución y acciones posteriores
        self._add_evolution_actions_section(doc)
        
        # Observaciones y cierre
        self._add_observations_closure_section(doc)
        
        self._add_professional_footer(doc)
        
        doc.save(output_path)
        return output_path
    
    def create_visit_report_template(self, output_path: str):
        """Crear plantilla mejorada de reporte de visita"""
        doc = Document()
        
        self._setup_document_margins(doc)
        self._add_professional_header(doc)
        
        self._add_main_title(doc, "REPORTE DE VISITA TÉCNICA")
        self._add_subtitle(doc, "INSPECCIÓN Y EVALUACIÓN DE INSTALACIONES")
        
        # Información de la obra
        self._add_work_info_section(doc)
        
        # Información del personal
        self._add_personnel_info_section(doc)
        
        # Roles y contactos
        self._add_roles_contacts_section(doc)
        
        # Uso de maquinaria
        self._add_machinery_usage_section(doc)
        
        # Observaciones por categoría
        self._add_observations_section(doc)
        
        # Firmas
        self._add_signatures_section(doc)
        
        self._add_professional_footer(doc)
        
        doc.save(output_path)
        return output_path
    
    def _setup_document_margins(self, doc):
        """Configurar márgenes profesionales del documento"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
    
    def _add_professional_header(self, doc):
        """Agregar header profesional con logo de Polifusión"""
        # Crear tabla para el header
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        
        # Configurar columnas
        logo_cell = header_table.cell(0, 0)
        info_cell = header_table.cell(0, 1)
        
        # Logo de Polifusión (texto estilizado)
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nombre de la empresa con estilo
        company_run = logo_para.add_run('POLIFUSIÓN')
        company_run.font.name = 'Arial'
        company_run.font.size = Pt(32)
        company_run.font.bold = True
        company_run.font.color.rgb = self.primary_color
        
        # Símbolo decorativo
        symbol_para = logo_cell.add_paragraph()
        symbol_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        symbol_run = symbol_para.add_run('⚡')
        symbol_run.font.size = Pt(28)
        symbol_run.font.color.rgb = self.accent_color
        
        # Información de la empresa
        info_para = info_cell.paragraphs[0]
        info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Tagline
        tagline_run = info_para.add_run('Especialistas en Tuberías y Fittings\n')
        tagline_run.font.name = 'Arial'
        tagline_run.font.size = Pt(14)
        tagline_run.font.bold = True
        tagline_run.font.color.rgb = self.secondary_color
        
        # Departamento
        dept_run = info_para.add_run('Departamento de Control de Calidad\n')
        dept_run.font.name = 'Arial'
        dept_run.font.size = Pt(12)
        dept_run.font.bold = True
        dept_run.font.color.rgb = self.text_color
        
        # Form ID
        form_run = info_para.add_run('FORM-SERTEC-0007\n')
        form_run.font.name = 'Arial'
        form_run.font.size = Pt(10)
        form_run.font.color.rgb = self.light_gray
        
        # Fecha de uso
        date_run = info_para.add_run('En uso desde: 28-08-2013')
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = self.light_gray
        
        # Línea separadora decorativa
        self._add_decorative_separator(doc)
    
    def _add_main_title(self, doc, title_text):
        """Agregar título principal del documento"""
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = self.primary_color
        
        # Espaciado después del título
        doc.add_paragraph()
    
    def _add_subtitle(self, doc, subtitle_text):
        """Agregar subtítulo del documento"""
        subtitle = doc.add_heading(subtitle_text, 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.bold = True
        subtitle_run.font.color.rgb = self.secondary_color
        
        # Espaciado después del subtítulo
        doc.add_paragraph()
    
    def _add_decorative_separator(self, doc):
        """Agregar separador decorativo"""
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_run = separator.add_run('═' * 80)
        separator_run.font.name = 'Arial'
        separator_run.font.size = Pt(12)
        separator_run.font.color.rgb = self.secondary_color
        
        # Espaciado
        doc.add_paragraph()
    
    def _add_applicant_info_section(self, doc):
        """Agregar sección de información del solicitante"""
        # Título de sección
        self._add_section_title(doc, "INFORMACIÓN DEL SOLICITANTE")
        
        # Tabla de información
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Configurar columnas
        cols = table.columns
        cols[0].width = Inches(2.5)
        cols[1].width = Inches(4.5)
        
        # Datos del solicitante
        data = [
            ("Solicitante:", "{{SOLICITANTE}}"),
            ("Fecha de Solicitud:", "{{FECHA_SOLICITUD}}"),
            ("Cliente:", "{{CLIENTE}}"),
            ("Informante:", "{{INFORMANTE}}")
        ]
        
        for i, (label, value) in enumerate(data):
            # Etiqueta
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(11)
            label_run.font.bold = True
            label_run.font.color.rgb = self.text_color
            
            # Valor
            value_cell = table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.name = 'Arial'
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_sample_description_section(self, doc):
        """Agregar sección de descripción de la muestra"""
        self._add_section_title(doc, "DESCRIPCIÓN DE LA MUESTRA")
        
        # Tabla de información técnica
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        cols = table.columns
        cols[0].width = Inches(2.5)
        cols[1].width = Inches(4.5)
        
        data = [
            ("Diámetro:", "{{DIAMETRO}} mm"),
            ("Proyecto:", "{{PROYECTO}}"),
            ("Ubicación:", "{{UBICACION}}"),
            ("Presión:", "{{PRESION}} bar"),
            ("Temperatura:", "{{TEMPERATURA}}")
        ]
        
        for i, (label, value) in enumerate(data):
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(11)
            label_run.font.bold = True
            label_run.font.color.rgb = self.text_color
            
            value_cell = table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.name = 'Arial'
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_tests_section(self, doc):
        """Agregar sección de ensayos realizados"""
        self._add_section_title(doc, "ENSAYOS REALIZADOS")
        
        # Lista de ensayos
        para = doc.add_paragraph()
        para_run = para.add_run("• Análisis visual de la muestra")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        
        para = doc.add_paragraph()
        para_run = para.add_run("• Evaluación de la fusión")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        
        para = doc.add_paragraph()
        para_run = para.add_run("• Análisis de fractura y cristalización")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        
        para = doc.add_paragraph()
        para_run = para.add_run("• {{ENSAYOS_ADICIONALES}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_detailed_comments_section(self, doc):
        """Agregar sección de comentarios detallados"""
        self._add_section_title(doc, "COMENTARIOS DETALLADOS")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{COMENTARIOS_DETALLADOS}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_conclusions_section(self, doc):
        """Agregar sección de conclusiones"""
        self._add_section_title(doc, "CONCLUSIONES")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{CONCLUSIONES_DETALLADAS}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_expert_signature_section(self, doc):
        """Agregar sección de firma del experto"""
        self._add_section_title(doc, "FIRMA DEL EXPERTO")
        
        # Tabla para firma
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        cols = table.columns
        cols[0].width = Inches(3)
        cols[1].width = Inches(4)
        
        # Nombre del experto
        name_cell = table.cell(0, 0)
        name_para = name_cell.paragraphs[0]
        name_run = name_para.add_run("Nombre del Experto:")
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(11)
        name_run.font.bold = True
        
        expert_cell = table.cell(0, 1)
        expert_para = expert_cell.paragraphs[0]
        expert_run = expert_para.add_run("{{EXPERTO_NOMBRE}}")
        expert_run.font.name = 'Arial'
        expert_run.font.size = Pt(11)
        expert_run.font.color.rgb = self.primary_color
        
        # Firma
        signature_cell = table.cell(1, 0)
        signature_para = signature_cell.paragraphs[0]
        signature_run = signature_para.add_run("Firma:")
        signature_run.font.name = 'Arial'
        signature_run.font.size = Pt(11)
        signature_run.font.bold = True
        
        # Espacio para firma
        signature_space_cell = table.cell(1, 1)
        signature_space_para = signature_space_cell.paragraphs[0]
        signature_space_para.add_run("_________________________")
        
        doc.add_paragraph()
    
    def _add_section_title(self, doc, title_text):
        """Agregar título de sección"""
        title = doc.add_heading(title_text, 2)
        title_run = title.runs[0]
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(13)
        title_run.font.bold = True
        title_run.font.color.rgb = self.primary_color
        
        # Línea decorativa
        line = doc.add_paragraph()
        line_run = line.add_run("─" * 50)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(8)
        line_run.font.color.rgb = self.secondary_color
    
    def _add_professional_footer(self, doc):
        """Agregar footer profesional"""
        # Línea separadora
        self._add_decorative_separator(doc)
        
        # Footer con información de la empresa
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_run = footer.add_run(f"{self.company_name} - {self.company_tagline}")
        footer_run.font.name = 'Arial'
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = self.light_gray
        
        # Fecha de generación
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run("Documento generado el: {{FECHA_GENERACION}}")
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(8)
        date_run.font.color.rgb = self.light_gray
    
    # Métodos para informe de incidencia
    def _add_incident_info_section(self, doc):
        """Agregar sección de información de la incidencia"""
        self._add_section_title(doc, "INFORMACIÓN DE LA INCIDENCIA")
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        cols = table.columns
        cols[0].width = Inches(2.5)
        cols[1].width = Inches(4.5)
        
        data = [
            ("Proveedor:", "{{PROVEEDOR}}"),
            ("Obra:", "{{OBRA}}"),
            ("Cliente:", "{{CLIENTE}}"),
            ("Fecha de Incidencia:", "{{FECHA_INCIDENCIA}}")
        ]
        
        for i, (label, value) in enumerate(data):
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(11)
            label_run.font.bold = True
            
            value_cell = table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.name = 'Arial'
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_problem_description_section(self, doc):
        """Agregar sección de descripción del problema"""
        self._add_section_title(doc, "DESCRIPCIÓN DEL PROBLEMA")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{DESCRIPCION_PROBLEMA}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_immediate_actions_section(self, doc):
        """Agregar sección de acciones inmediatas"""
        self._add_section_title(doc, "ACCIONES INMEDIATAS")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{ACCIONES_INMEDIATAS}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_evolution_actions_section(self, doc):
        """Agregar sección de evolución y acciones posteriores"""
        self._add_section_title(doc, "EVOLUCIÓN Y ACCIONES POSTERIORES")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{EVOLUCION_ACCIONES}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_observations_closure_section(self, doc):
        """Agregar sección de observaciones y cierre"""
        self._add_section_title(doc, "OBSERVACIONES Y CIERRE")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{OBSERVACIONES_CIERRE}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    # Métodos para reporte de visita
    def _add_work_info_section(self, doc):
        """Agregar sección de información de la obra"""
        self._add_section_title(doc, "INFORMACIÓN DE LA OBRA")
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        cols = table.columns
        cols[0].width = Inches(2.5)
        cols[1].width = Inches(4.5)
        
        data = [
            ("Obra:", "{{OBRA}}"),
            ("Cliente:", "{{CLIENTE}}"),
            ("Ubicación:", "{{UBICACION}}")
        ]
        
        for i, (label, value) in enumerate(data):
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(11)
            label_run.font.bold = True
            
            value_cell = table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.name = 'Arial'
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_personnel_info_section(self, doc):
        """Agregar sección de información del personal"""
        self._add_section_title(doc, "INFORMACIÓN DEL PERSONAL")
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        cols = table.columns
        cols[0].width = Inches(2.5)
        cols[1].width = Inches(4.5)
        
        data = [
            ("Vendedor:", "{{VENDEDOR}}"),
            ("Técnico:", "{{TECNICO}}"),
            ("Fecha de Visita:", "{{FECHA_VISITA}}")
        ]
        
        for i, (label, value) in enumerate(data):
            label_cell = table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(11)
            label_run.font.bold = True
            
            value_cell = table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.name = 'Arial'
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_roles_contacts_section(self, doc):
        """Agregar sección de roles y contactos"""
        self._add_section_title(doc, "ROLES Y CONTACTOS")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{ROLES_CONTACTOS}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_machinery_usage_section(self, doc):
        """Agregar sección de uso de maquinaria"""
        self._add_section_title(doc, "USO DE MAQUINARIA")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{USO_MAQUINARIA}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_observations_section(self, doc):
        """Agregar sección de observaciones por categoría"""
        self._add_section_title(doc, "OBSERVACIONES POR CATEGORÍA")
        
        para = doc.add_paragraph()
        para_run = para.add_run("{{OBSERVACIONES_CATEGORIA}}")
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(11)
        para_run.font.color.rgb = self.primary_color
        
        doc.add_paragraph()
    
    def _add_signatures_section(self, doc):
        """Agregar sección de firmas"""
        self._add_section_title(doc, "FIRMAS")
        
        # Tabla para firmas
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        cols = table.columns
        cols[0].width = Inches(3)
        cols[1].width = Inches(4)
        
        # Firma del técnico
        tech_cell = table.cell(0, 0)
        tech_para = tech_cell.paragraphs[0]
        tech_run = tech_para.add_run("Técnico:")
        tech_run.font.name = 'Arial'
        tech_run.font.size = Pt(11)
        tech_run.font.bold = True
        
        tech_signature_cell = table.cell(0, 1)
        tech_signature_para = tech_signature_cell.paragraphs[0]
        tech_signature_para.add_run("_________________________")
        
        # Firma del cliente
        client_cell = table.cell(1, 0)
        client_para = client_cell.paragraphs[0]
        client_run = client_para.add_run("Cliente:")
        client_run.font.name = 'Arial'
        client_run.font.size = Pt(11)
        client_run.font.bold = True
        
        client_signature_cell = table.cell(1, 1)
        client_signature_para = client_signature_cell.paragraphs[0]
        client_signature_para.add_run("_________________________")
        
        doc.add_paragraph()


# Instancia global del generador mejorado
polifusion_enhanced_generator = PolifusionEnhancedTemplateGenerator()
