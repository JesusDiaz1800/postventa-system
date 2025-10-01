"""
Generador de plantillas ultra profesionales para Polifusión
Diseño corporativo de alta calidad con logo real integrado
"""

import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class UltraProfessionalTemplateGenerator:
    """Generador de plantillas ultra profesionales para Polifusión"""
    
    def __init__(self):
        self.company_name = "POLIFUSIÓN"
        self.company_full_name = "POLIFUSIÓN S.A."
        self.company_tagline = "Especialistas en Tuberías y Fittings de Polipropileno"
        self.department = "Departamento de Control de Calidad"
        self.address = "Av. Libertador Bernardo O'Higgins 4687, Santiago, Chile"
        self.phone = "+56 2 2345 6789"
        self.email = "calidad@polifusion.cl"
        self.website = "www.polifusion.cl"
        
        # Colores corporativos ultra profesionales
        self.primary_blue = RGBColor(0, 51, 102)      # Azul corporativo profundo
        self.secondary_blue = RGBColor(0, 102, 204)   # Azul medio
        self.accent_orange = RGBColor(255, 140, 0)    # Naranja corporativo
        self.dark_gray = RGBColor(45, 45, 45)         # Gris muy oscuro
        self.medium_gray = RGBColor(100, 100, 100)    # Gris medio
        self.light_gray = RGBColor(240, 240, 240)     # Gris claro
        self.white = RGBColor(255, 255, 255)          # Blanco
        self.success_green = RGBColor(0, 128, 0)      # Verde éxito
        
        # Fuentes corporativas ultra profesionales
        self.title_font_size = Pt(20)
        self.subtitle_font_size = Pt(16)
        self.header_font_size = Pt(14)
        self.body_font_size = Pt(12)
        self.small_font_size = Pt(10)
        self.tiny_font_size = Pt(8)
    
    def create_ultra_professional_lab_report(self, output_path: str):
        """Crear informe de laboratorio ultra profesional"""
        doc = Document()
        
        # Configurar página profesional
        self._setup_ultra_professional_page(doc)
        
        # Header corporativo ultra profesional
        self._add_ultra_professional_header(doc)
        
        # Título del documento con diseño premium
        self._add_premium_document_title(doc, "INFORME DE LABORATORIO")
        self._add_premium_document_subtitle(doc, "ANÁLISIS TÉCNICO DE TUBERÍAS PP-RCT")
        
        # Información del solicitante con diseño premium
        self._add_premium_applicant_section(doc)
        
        # Información técnica con diseño premium
        self._add_premium_technical_info_section(doc)
        
        # Ensayos realizados con diseño premium
        self._add_premium_tests_section(doc)
        
        # Análisis detallado con diseño premium
        self._add_premium_analysis_section(doc)
        
        # Conclusiones con diseño premium
        self._add_premium_conclusions_section(doc)
        
        # Firmas con diseño premium
        self._add_premium_signatures_section(doc)
        
        # Footer corporativo ultra profesional
        self._add_ultra_professional_footer(doc)
        
        # Guardar documento
        doc.save(output_path)
        print(f"✅ Informe de laboratorio ultra profesional creado: {output_path}")
    
    def create_ultra_professional_incident_report(self, output_path: str):
        """Crear informe de incidencia ultra profesional"""
        doc = Document()
        
        # Configurar página profesional
        self._setup_ultra_professional_page(doc)
        
        # Header corporativo ultra profesional
        self._add_ultra_professional_header(doc)
        
        # Título del documento con diseño premium
        self._add_premium_document_title(doc, "INFORME DE INCIDENCIA")
        self._add_premium_document_subtitle(doc, "REGISTRO DE PROBLEMAS Y ACCIONES CORRECTIVAS")
        
        # Información de la incidencia con diseño premium
        self._add_premium_incident_info_section(doc)
        
        # Descripción del problema con diseño premium
        self._add_premium_problem_description_section(doc)
        
        # Acciones tomadas con diseño premium
        self._add_premium_actions_section(doc)
        
        # Seguimiento con diseño premium
        self._add_premium_follow_up_section(doc)
        
        # Cierre con diseño premium
        self._add_premium_closure_section(doc)
        
        # Firmas con diseño premium
        self._add_premium_signatures_section(doc)
        
        # Footer corporativo ultra profesional
        self._add_ultra_professional_footer(doc)
        
        # Guardar documento
        doc.save(output_path)
        print(f"✅ Informe de incidencia ultra profesional creado: {output_path}")
    
    def create_ultra_professional_visit_report(self, output_path: str):
        """Crear informe de visita ultra profesional"""
        doc = Document()
        
        # Configurar página profesional
        self._setup_ultra_professional_page(doc)
        
        # Header corporativo ultra profesional
        self._add_ultra_professional_header(doc)
        
        # Título del documento con diseño premium
        self._add_premium_document_title(doc, "REPORTE DE VISITA TÉCNICA")
        self._add_premium_document_subtitle(doc, "EVALUACIÓN DE INSTALACIÓN Y SERVICIO")
        
        # Información de la visita con diseño premium
        self._add_premium_visit_info_section(doc)
        
        # Personal presente con diseño premium
        self._add_premium_personnel_section(doc)
        
        # Evaluación técnica con diseño premium
        self._add_premium_technical_evaluation_section(doc)
        
        # Observaciones con diseño premium
        self._add_premium_observations_section(doc)
        
        # Recomendaciones con diseño premium
        self._add_premium_recommendations_section(doc)
        
        # Firmas con diseño premium
        self._add_premium_signatures_section(doc)
        
        # Footer corporativo ultra profesional
        self._add_ultra_professional_footer(doc)
        
        # Guardar documento
        doc.save(output_path)
        print(f"✅ Reporte de visita ultra profesional creado: {output_path}")
    
    def _setup_ultra_professional_page(self, doc):
        """Configurar página ultra profesional"""
        section = doc.sections[0]
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)    # A4
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
    
    def _add_ultra_professional_header(self, doc):
        """Agregar header corporativo ultra profesional"""
        # Crear tabla para header con diseño premium
        header_table = doc.add_table(rows=1, cols=3)
        header_table.width = Inches(6.5)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar columnas
        header_table.columns[0].width = Inches(1.8)
        header_table.columns[1].width = Inches(3.2)
        header_table.columns[2].width = Inches(1.5)
        
        # Celda 1: Logo real (placeholder para imagen)
        logo_cell = header_table.cell(0, 0)
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Intentar agregar logo real
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', '..', 'static', 'images', 'polifusion_logo.png')
        if os.path.exists(logo_path):
            try:
                logo_para.add_run().add_picture(logo_path, width=Inches(1.5))
            except:
                # Si falla, usar texto
                logo_run = logo_para.add_run("POLIFUSIÓN")
                logo_run.font.size = Pt(18)
                logo_run.font.color.rgb = self.primary_blue
                logo_run.bold = True
        else:
            # Logo de texto profesional
            logo_run = logo_para.add_run("POLIFUSIÓN")
            logo_run.font.size = Pt(18)
            logo_run.font.color.rgb = self.primary_blue
            logo_run.bold = True
        
        # Celda 2: Información de la empresa con diseño premium
        info_cell = header_table.cell(0, 1)
        info_para = info_cell.paragraphs[0]
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nombre de la empresa
        company_run = info_para.add_run(self.company_full_name)
        company_run.font.size = Pt(16)
        company_run.font.color.rgb = self.primary_blue
        company_run.bold = True
        
        # Tagline
        tagline_run = info_para.add_run(f"\n{self.company_tagline}")
        tagline_run.font.size = Pt(11)
        tagline_run.font.color.rgb = self.medium_gray
        tagline_run.italic = True
        
        # Departamento
        dept_run = info_para.add_run(f"\n{self.department}")
        dept_run.font.size = Pt(10)
        dept_run.font.color.rgb = self.medium_gray
        dept_run.bold = True
        
        # Celda 3: Información de contacto con diseño premium
        contact_cell = header_table.cell(0, 2)
        contact_para = contact_cell.paragraphs[0]
        contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Dirección
        addr_run = contact_para.add_run(self.address)
        addr_run.font.size = Pt(9)
        addr_run.font.color.rgb = self.medium_gray
        
        # Teléfono
        phone_run = contact_para.add_run(f"\n{self.phone}")
        phone_run.font.size = Pt(9)
        phone_run.font.color.rgb = self.medium_gray
        
        # Email
        email_run = contact_para.add_run(f"\n{self.email}")
        email_run.font.size = Pt(9)
        email_run.font.color.rgb = self.medium_gray
        
        # Línea separadora premium
        self._add_premium_separator_line(doc)
    
    def _add_premium_document_title(self, doc, title):
        """Agregar título del documento con diseño premium"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(8)
        
        title_run = para.add_run(title)
        title_run.font.size = self.title_font_size
        title_run.font.color.rgb = self.primary_blue
        title_run.bold = True
        
        # Línea decorativa bajo el título
        decor_para = doc.add_paragraph()
        decor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        decor_para.space_after = Pt(12)
        
        decor_run = decor_para.add_run("◆" * 20)
        decor_run.font.size = Pt(12)
        decor_run.font.color.rgb = self.accent_orange
    
    def _add_premium_document_subtitle(self, doc, subtitle):
        """Agregar subtítulo del documento con diseño premium"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(16)
        
        subtitle_run = para.add_run(subtitle)
        subtitle_run.font.size = self.subtitle_font_size
        subtitle_run.font.color.rgb = self.dark_gray
        subtitle_run.italic = True
    
    def _add_premium_separator_line(self, doc):
        """Agregar línea separadora premium"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(16)
        
        line_run = para.add_run("▬" * 50)
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = self.light_gray
    
    def _add_premium_applicant_section(self, doc):
        """Agregar sección de información del solicitante con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "INFORMACIÓN DEL SOLICITANTE")
        
        # Tabla de información con diseño premium
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.width = Inches(6.5)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar columnas
        info_table.columns[0].width = Inches(2.2)
        info_table.columns[1].width = Inches(4.3)
        
        # Datos del solicitante
        applicant_data = [
            ("Solicitante:", "{{SOLICITANTE}}"),
            ("Fecha de Solicitud:", "{{FECHA_SOLICITUD}}"),
            ("Cliente:", "{{CLIENTE}}"),
            ("Informante:", "{{INFORMANTE}}")
        ]
        
        for i, (label, value) in enumerate(applicant_data):
            # Etiqueta con diseño premium
            label_cell = info_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.size = self.body_font_size
            label_run.font.color.rgb = self.primary_blue
            label_run.bold = True
            
            # Valor con diseño premium
            value_cell = info_table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.size = self.body_font_size
            value_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_technical_info_section(self, doc):
        """Agregar sección de información técnica con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "INFORMACIÓN TÉCNICA")
        
        # Tabla de información técnica con diseño premium
        tech_table = doc.add_table(rows=5, cols=2)
        tech_table.style = 'Table Grid'
        tech_table.width = Inches(6.5)
        tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar columnas
        tech_table.columns[0].width = Inches(2.2)
        tech_table.columns[1].width = Inches(4.3)
        
        # Datos técnicos
        tech_data = [
            ("Diámetro:", "{{DIAMETRO}} mm"),
            ("Proyecto:", "{{PROYECTO}}"),
            ("Ubicación:", "{{UBICACION}}"),
            ("Presión:", "{{PRESION}} bar"),
            ("Temperatura:", "{{TEMPERATURA}} °C")
        ]
        
        for i, (label, value) in enumerate(tech_data):
            # Etiqueta con diseño premium
            label_cell = tech_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.size = self.body_font_size
            label_run.font.color.rgb = self.primary_blue
            label_run.bold = True
            
            # Valor con diseño premium
            value_cell = tech_table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.size = self.body_font_size
            value_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_tests_section(self, doc):
        """Agregar sección de ensayos con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "ENSAYOS REALIZADOS")
        
        # Contenido de ensayos con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        tests_run = para.add_run("{{ENSAYOS_ADICIONALES}}")
        tests_run.font.size = self.body_font_size
        tests_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_analysis_section(self, doc):
        """Agregar sección de análisis con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "ANÁLISIS DETALLADO")
        
        # Contenido del análisis con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        analysis_run = para.add_run("{{COMENTARIOS_DETALLADOS}}")
        analysis_run.font.size = self.body_font_size
        analysis_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_conclusions_section(self, doc):
        """Agregar sección de conclusiones con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "CONCLUSIONES")
        
        # Contenido de conclusiones con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        conclusions_run = para.add_run("{{CONCLUSIONES_DETALLADAS}}")
        conclusions_run.font.size = self.body_font_size
        conclusions_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_incident_info_section(self, doc):
        """Agregar sección de información de incidencia con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "INFORMACIÓN DE LA INCIDENCIA")
        
        # Tabla de información con diseño premium
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.width = Inches(6.5)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar columnas
        info_table.columns[0].width = Inches(2.2)
        info_table.columns[1].width = Inches(4.3)
        
        # Datos de la incidencia
        incident_data = [
            ("Proveedor:", "{{PROVEEDOR}}"),
            ("Obra:", "{{OBRA}}"),
            ("Cliente:", "{{CLIENTE}}"),
            ("Fecha de Detección:", "{{FECHA_INCIDENCIA}}")
        ]
        
        for i, (label, value) in enumerate(incident_data):
            # Etiqueta con diseño premium
            label_cell = info_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.size = self.body_font_size
            label_run.font.color.rgb = self.primary_blue
            label_run.bold = True
            
            # Valor con diseño premium
            value_cell = info_table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.size = self.body_font_size
            value_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_problem_description_section(self, doc):
        """Agregar sección de descripción del problema con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "DESCRIPCIÓN DEL PROBLEMA")
        
        # Contenido del problema con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        problem_run = para.add_run("{{DESCRIPCION_PROBLEMA}}")
        problem_run.font.size = self.body_font_size
        problem_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_actions_section(self, doc):
        """Agregar sección de acciones con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "ACCIONES INMEDIATAS")
        
        # Contenido de acciones con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        actions_run = para.add_run("{{ACCIONES_INMEDIATAS}}")
        actions_run.font.size = self.body_font_size
        actions_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
        
        # Evolución
        self._add_premium_section_title(doc, "EVOLUCIÓN Y ACCIONES POSTERIORES")
        
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        evolution_run = para.add_run("{{EVOLUCION_ACCIONES}}")
        evolution_run.font.size = self.body_font_size
        evolution_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_follow_up_section(self, doc):
        """Agregar sección de seguimiento con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "SEGUIMIENTO")
        
        # Contenido de seguimiento con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        follow_run = para.add_run("{{OBSERVACIONES}}")
        follow_run.font.size = self.body_font_size
        follow_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_closure_section(self, doc):
        """Agregar sección de cierre con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "CIERRE DE LA INCIDENCIA")
        
        # Contenido de cierre con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        closure_run = para.add_run("{{OBSERVACIONES_CIERRE}}")
        closure_run.font.size = self.body_font_size
        closure_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_visit_info_section(self, doc):
        """Agregar sección de información de visita con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "INFORMACIÓN DE LA VISITA")
        
        # Tabla de información con diseño premium
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.width = Inches(6.5)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar columnas
        info_table.columns[0].width = Inches(2.2)
        info_table.columns[1].width = Inches(4.3)
        
        # Datos de la visita
        visit_data = [
            ("Obra:", "{{OBRA}}"),
            ("Cliente:", "{{CLIENTE}}"),
            ("Vendedor:", "{{VENDEDOR}}"),
            ("Técnico:", "{{TECNICO}}")
        ]
        
        for i, (label, value) in enumerate(visit_data):
            # Etiqueta con diseño premium
            label_cell = info_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.size = self.body_font_size
            label_run.font.color.rgb = self.primary_blue
            label_run.bold = True
            
            # Valor con diseño premium
            value_cell = info_table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.size = self.body_font_size
            value_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_personnel_section(self, doc):
        """Agregar sección de personal con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "PERSONAL PRESENTE")
        
        # Contenido del personal con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        personnel_run = para.add_run("{{PERSONAL_INFO}}")
        personnel_run.font.size = self.body_font_size
        personnel_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
        
        # Roles y contactos
        self._add_premium_section_title(doc, "ROLES Y CONTACTOS")
        
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        roles_run = para.add_run("{{ROLES_CONTACTOS}}")
        roles_run.font.size = self.body_font_size
        roles_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_technical_evaluation_section(self, doc):
        """Agregar sección de evaluación técnica con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "EVALUACIÓN TÉCNICA")
        
        # Uso de maquinaria con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        machinery_run = para.add_run("Uso de Maquinaria:\n{{MAQUINARIA_USO}}")
        machinery_run.font.size = self.body_font_size
        machinery_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_observations_section(self, doc):
        """Agregar sección de observaciones con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "OBSERVACIONES")
        
        # Observaciones por categoría con diseño premium
        observations = [
            ("Instalación:", "{{OBSERVACIONES_INSTALACION}}"),
            ("Material:", "{{OBSERVACIONES_MATERIAL}}"),
            ("Técnico:", "{{OBSERVACIONES_TECNICO}}"),
            ("General:", "{{OBSERVACIONES_GENERAL}}")
        ]
        
        for category, content in observations:
            para = doc.add_paragraph()
            para.space_after = Pt(6)
            
            cat_run = para.add_run(f"{category}\n{content}")
            cat_run.font.size = self.body_font_size
            cat_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_recommendations_section(self, doc):
        """Agregar sección de recomendaciones con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "RECOMENDACIONES")
        
        # Contenido de recomendaciones con diseño premium
        para = doc.add_paragraph()
        para.space_after = Pt(8)
        
        rec_run = para.add_run("{{RECOMENDACIONES}}")
        rec_run.font.size = self.body_font_size
        rec_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_signatures_section(self, doc):
        """Agregar sección de firmas con diseño premium"""
        # Título de sección premium
        self._add_premium_section_title(doc, "FIRMAS")
        
        # Tabla de firmas con diseño premium
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.width = Inches(6.5)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar columnas
        sig_table.columns[0].width = Inches(3.25)
        sig_table.columns[1].width = Inches(3.25)
        
        # Firmas
        signatures = [
            ("Responsable Técnico:", "{{EXPERTO_NOMBRE}}"),
            ("Fecha:", "{{FECHA_GENERACION}}")
        ]
        
        for i, (label, value) in enumerate(signatures):
            # Etiqueta con diseño premium
            label_cell = sig_table.cell(i, 0)
            label_para = label_cell.paragraphs[0]
            label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = label_para.add_run(label)
            label_run.font.size = self.body_font_size
            label_run.font.color.rgb = self.primary_blue
            label_run.bold = True
            
            # Valor con diseño premium
            value_cell = sig_table.cell(i, 1)
            value_para = value_cell.paragraphs[0]
            value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            value_run = value_para.add_run(value)
            value_run.font.size = self.body_font_size
            value_run.font.color.rgb = self.dark_gray
        
        doc.add_paragraph()  # Espacio
    
    def _add_premium_section_title(self, doc, title):
        """Agregar título de sección premium"""
        para = doc.add_paragraph()
        para.space_before = Pt(16)
        para.space_after = Pt(8)
        
        title_run = para.add_run(title)
        title_run.font.size = self.header_font_size
        title_run.font.color.rgb = self.primary_blue
        title_run.bold = True
        
        # Línea decorativa bajo el título
        decor_para = doc.add_paragraph()
        decor_para.space_after = Pt(8)
        
        decor_run = decor_para.add_run("▬" * 30)
        decor_run.font.size = Pt(8)
        decor_run.font.color.rgb = self.accent_orange
    
    def _add_ultra_professional_footer(self, doc):
        """Agregar footer corporativo ultra profesional"""
        # Línea separadora premium
        self._add_premium_separator_line(doc)
        
        # Footer con diseño premium
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.space_after = Pt(6)
        
        footer_run = footer_para.add_run(f"{self.company_full_name} | {self.address} | {self.phone} | {self.email}")
        footer_run.font.size = self.small_font_size
        footer_run.font.color.rgb = self.medium_gray
