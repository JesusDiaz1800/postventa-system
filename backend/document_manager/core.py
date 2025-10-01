"""
Módulo principal de gestión de documentos
Integra WeasyPrint, Jinja2 y python-docx para generación de documentos profesionales
"""

import os
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from jinja2 import Environment, FileSystemLoader, Template
import weasyprint
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration
from PIL import Image
import io
import base64

from .templates import TemplateManager
from .ai_processor import AIProcessor
from .file_manager import FileManager

logger = logging.getLogger(__name__)

class DocumentManager:
    """
    Gestor principal de documentos para el sistema de control de calidad
    """
    
    def __init__(self, base_path: str = None):
        """
        Inicializar el gestor de documentos
        
        Args:
            base_path: Ruta base para almacenar documentos
        """
        self.base_path = base_path or os.path.join(os.getcwd(), 'documents')
        self.template_manager = TemplateManager()
        self.ai_processor = AIProcessor()
        self.file_manager = FileManager(self.base_path)
        
        # Crear directorios necesarios
        self._create_directories()
        
        # Configurar Jinja2
        self.jinja_env = Environment(
            loader=FileSystemLoader(self.template_manager.templates_path),
            autoescape=True
        )
        
        logger.info(f"DocumentManager inicializado en: {self.base_path}")
    
    def _create_directories(self):
        """Crear estructura de directorios necesaria"""
        directories = [
            'templates',
            'output',
            'temp',
            'library',
            'shared'
        ]
        
        for directory in directories:
            path = os.path.join(self.base_path, directory)
            os.makedirs(path, exist_ok=True)
    
    def generar_docx(self, contexto: Dict[str, Any], template_name: str = "incident_report") -> str:
        """
        Generar documento Word desde plantilla y datos dinámicos
        
        Args:
            contexto: Diccionario con datos para la plantilla
            template_name: Nombre de la plantilla a usar
            
        Returns:
            Ruta del archivo Word generado
        """
        try:
            # Obtener plantilla
            template_path = self.template_manager.get_template(template_name, "docx")
            
            # Crear documento Word
            doc = Document(template_path)
            
            # Procesar contexto con IA para mejorar redacción
            contexto_mejorado = self.ai_processor.maquillar_redaccion(contexto)
            
            # Aplicar datos al documento
            self._apply_context_to_docx(doc, contexto_mejorado)
            
            # Generar nombre de archivo único
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{template_name}_{timestamp}.docx"
            output_path = os.path.join(self.base_path, 'output', filename)
            
            # Guardar documento
            doc.save(output_path)
            
            logger.info(f"Documento Word generado: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generando documento Word: {str(e)}")
            raise
    
    def _apply_context_to_docx(self, doc: Document, contexto: Dict[str, Any]):
        """
        Aplicar contexto a documento Word
        
        Args:
            doc: Documento Word
            contexto: Datos a aplicar
        """
        # Reemplazar placeholders en párrafos
        for paragraph in doc.paragraphs:
            for key, value in contexto.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
        
        # Reemplazar placeholders en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in contexto.items():
                            if f"{{{{{key}}}}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
    
    def editar_docx(self, path: str) -> bool:
        """
        Permitir al usuario editar el documento Word
        
        Args:
            path: Ruta del archivo Word
            
        Returns:
            True si se abrió correctamente
        """
        try:
            import subprocess
            import platform
            
            if platform.system() == "Windows":
                os.startfile(path)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", path])
            else:  # Linux
                subprocess.run(["xdg-open", path])
            
            logger.info(f"Documento Word abierto para edición: {path}")
            return True
            
        except Exception as e:
            logger.error(f"Error abriendo documento Word: {str(e)}")
            return False
    
    def guardar_pdf(self, path_docx: str, path_pdf: str = None) -> str:
        """
        Convertir documento Word a PDF profesional con WeasyPrint
        
        Args:
            path_docx: Ruta del archivo Word
            path_pdf: Ruta de salida del PDF (opcional)
            
        Returns:
            Ruta del archivo PDF generado
        """
        try:
            # Generar HTML desde Word
            html_content = self._docx_to_html(path_docx)
            
            # Generar PDF con WeasyPrint
            if not path_pdf:
                path_pdf = path_docx.replace('.docx', '.pdf')
            
            # Configurar CSS para diseño profesional
            css_content = self._get_professional_css()
            
            # Generar PDF
            html_doc = HTML(string=html_content)
            css_doc = CSS(string=css_content)
            
            html_doc.write_pdf(path_pdf, stylesheets=[css_doc])
            
            logger.info(f"PDF generado: {path_pdf}")
            return path_pdf
            
        except Exception as e:
            logger.error(f"Error generando PDF: {str(e)}")
            raise
    
    def _docx_to_html(self, path_docx: str) -> str:
        """
        Convertir documento Word a HTML
        
        Args:
            path_docx: Ruta del archivo Word
            
        Returns:
            Contenido HTML
        """
        doc = Document(path_docx)
        html_content = []
        
        # Procesar párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                style = paragraph.style.name
                if 'Heading' in style:
                    level = style.replace('Heading ', '')
                    html_content.append(f"<h{level}>{paragraph.text}</h{level}>")
                else:
                    html_content.append(f"<p>{paragraph.text}</p>")
        
        # Procesar tablas
        for table in doc.tables:
            html_content.append("<table class='professional-table'>")
            for i, row in enumerate(table.rows):
                tag = 'th' if i == 0 else 'td'
                html_content.append("<tr>")
                for cell in row.cells:
                    html_content.append(f"<{tag}>{cell.text}</{tag}>")
                html_content.append("</tr>")
            html_content.append("</table>")
        
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Informe de Incidencia</title>
        </head>
        <body>
            {''.join(html_content)}
        </body>
        </html>
        """
    
    def _get_professional_css(self) -> str:
        """
        Obtener CSS profesional para los PDFs
        
        Returns:
            Contenido CSS
        """
        return """
        @page {
            size: A4;
            margin: 2cm;
            @top-center {
                content: "Polifusión S.A. - Sistema de Control de Calidad";
                font-size: 10pt;
                color: #126FCC;
            }
            @bottom-center {
                content: "Confidencial - Documento interno";
                font-size: 8pt;
                color: #666;
            }
        }
        
        body {
            font-family: 'Helvetica', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
        }
        
        h1 {
            color: #126FCC;
            font-size: 18pt;
            text-align: center;
            margin-bottom: 20pt;
            border-bottom: 2px solid #126FCC;
            padding-bottom: 10pt;
        }
        
        h2 {
            color: #126FCC;
            font-size: 14pt;
            margin-top: 20pt;
            margin-bottom: 10pt;
        }
        
        h3 {
            color: #333;
            font-size: 12pt;
            margin-top: 15pt;
            margin-bottom: 8pt;
        }
        
        .professional-table {
            width: 100%;
            border-collapse: collapse;
            margin: 15pt 0;
            font-size: 10pt;
        }
        
        .professional-table th {
            background-color: #126FCC;
            color: white;
            padding: 8pt;
            text-align: left;
            font-weight: bold;
        }
        
        .professional-table td {
            padding: 6pt 8pt;
            border-bottom: 1px solid #ddd;
        }
        
        .professional-table tr:nth-child(even) {
            background-color: #f8f9fa;
        }
        
        .header-info {
            background-color: #f0f8ff;
            padding: 15pt;
            border-left: 4px solid #126FCC;
            margin-bottom: 20pt;
        }
        
        .footer-info {
            margin-top: 30pt;
            padding-top: 15pt;
            border-top: 1px solid #ddd;
            font-size: 9pt;
            color: #666;
        }
        
        .image-container {
            text-align: center;
            margin: 15pt 0;
        }
        
        .image-container img {
            max-width: 100%;
            height: auto;
            border: 1px solid #ddd;
            border-radius: 4px;
        }
        
        .highlight {
            background-color: #fff3cd;
            padding: 10pt;
            border-left: 4px solid #ffc107;
            margin: 10pt 0;
        }
        
        .critical {
            background-color: #f8d7da;
            padding: 10pt;
            border-left: 4px solid #dc3545;
            margin: 10pt 0;
        }
        """
    
    def insertar_imagen(self, path_docx: str, imagen_path: str, descripcion: str = "") -> str:
        """
        Insertar imagen en documento Word
        
        Args:
            path_docx: Ruta del documento Word
            imagen_path: Ruta de la imagen
            descripcion: Descripción de la imagen
            
        Returns:
            Ruta del documento actualizado
        """
        try:
            doc = Document(path_docx)
            
            # Agregar imagen
            paragraph = doc.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # Insertar imagen
            run.add_picture(imagen_path, width=Inches(4))
            
            # Agregar descripción si se proporciona
            if descripcion:
                desc_paragraph = doc.add_paragraph()
                desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                desc_run = desc_paragraph.add_run(descripcion)
                desc_run.font.size = Pt(9)
                desc_run.font.italic = True
            
            # Guardar documento actualizado
            doc.save(path_docx)
            
            logger.info(f"Imagen insertada en documento: {path_docx}")
            return path_docx
            
        except Exception as e:
            logger.error(f"Error insertando imagen: {str(e)}")
            raise
    
    def analizar_imagenes(self, imagenes: List[str]) -> Dict[str, Any]:
        """
        Analizar imágenes con IA para encontrar posibles causas
        
        Args:
            imagenes: Lista de rutas de imágenes
            
        Returns:
            Diccionario con análisis de cada imagen
        """
        return self.ai_processor.analizar_imagenes(imagenes)
    
    def generar_documento_completo(self, contexto: Dict[str, Any], imagenes: List[str] = None) -> Dict[str, str]:
        """
        Generar documento completo (Word + PDF) con imágenes
        
        Args:
            contexto: Datos del contexto
            imagenes: Lista de rutas de imágenes
            
        Returns:
            Diccionario con rutas de archivos generados
        """
        try:
            # Generar documento Word
            docx_path = self.generar_docx(contexto)
            
            # Insertar imágenes si se proporcionan
            if imagenes:
                for imagen in imagenes:
                    self.insertar_imagen(docx_path, imagen)
            
            # Generar PDF
            pdf_path = self.guardar_pdf(docx_path)
            
            # Guardar en biblioteca
            self.file_manager.guardar_en_biblioteca(docx_path, pdf_path, contexto)
            
            return {
                'docx': docx_path,
                'pdf': pdf_path,
                'biblioteca_docx': self.file_manager.get_library_path(docx_path),
                'biblioteca_pdf': self.file_manager.get_library_path(pdf_path)
            }
            
        except Exception as e:
            logger.error(f"Error generando documento completo: {str(e)}")
            raise
    
    def obtener_biblioteca_documentos(self) -> List[Dict[str, Any]]:
        """
        Obtener lista de documentos en la biblioteca
        
        Returns:
            Lista de documentos con metadatos
        """
        return self.file_manager.obtener_biblioteca()
