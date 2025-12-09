#!/usr/bin/env python3
"""
Script para verificar la calidad de los documentos generados
"""

import os
import sys
import django
from pathlib import Path

# Setup Django
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'postventa_system.settings-sqlserver')
django.setup()

from django.conf import settings

def check_document_quality():
    """Verificar la calidad de los documentos generados"""
    
    print("🔍 VERIFICANDO CALIDAD DE DOCUMENTOS")
    print("=" * 60)
    
    documents_dir = os.path.join(settings.SHARED_FOLDER_PATH, 'documents')
    
    if not os.path.exists(documents_dir):
        print(f"❌ Directorio de documentos no existe: {documents_dir}")
        return
    
    print(f"📁 Directorio de documentos: {documents_dir}")
    
    # Listar archivos
    files = os.listdir(documents_dir)
    docx_files = [f for f in files if f.endswith('.docx')]
    
    print(f"\n📋 Documentos encontrados: {len(docx_files)}")
    
    for i, file in enumerate(docx_files[:3], 1):  # Revisar solo los primeros 3
        file_path = os.path.join(documents_dir, file)
        file_size = os.path.getsize(file_path)
        
        print(f"\n{i}. 📄 {file}")
        print(f"   📏 Tamaño: {file_size:,} bytes")
        print(f"   📅 Modificado: {os.path.getmtime(file_path)}")
        
        # Verificar si el archivo se puede abrir
        try:
            from docx import Document
            doc = Document(file_path)
            
            print(f"   ✅ Archivo válido (se puede abrir)")
            print(f"   📝 Párrafos: {len(doc.paragraphs)}")
            print(f"   📊 Tablas: {len(doc.tables)}")
            
            # Verificar contenido
            has_content = False
            for para in doc.paragraphs[:5]:  # Primeros 5 párrafos
                if para.text.strip():
                    print(f"   📄 Contenido: {para.text[:50]}...")
                    has_content = True
                    break
            
            if not has_content:
                print(f"   ⚠️  Sin contenido visible")
                
        except Exception as e:
            print(f"   ❌ Error al abrir: {e}")
    
    # Verificar plantillas
    print(f"\n🎨 VERIFICANDO PLANTILLAS:")
    templates_dir = os.path.join(settings.SHARED_FOLDER_PATH, 'templates')
    
    if os.path.exists(templates_dir):
        template_files = [f for f in os.listdir(templates_dir) if f.endswith('.docx')]
        print(f"   📁 Plantillas encontradas: {len(template_files)}")
        
        for template in template_files:
            template_path = os.path.join(templates_dir, template)
            template_size = os.path.getsize(template_path)
            print(f"   📄 {template} ({template_size:,} bytes)")
    else:
        print(f"   ❌ Directorio de plantillas no existe: {templates_dir}")
    
    print("\n🎉 ¡VERIFICACIÓN COMPLETADA!")

if __name__ == "__main__":
    try:
        check_document_quality()
    except Exception as e:
        print(f"❌ Error durante la verificación: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
